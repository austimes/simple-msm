"""
Generate a combined non-road transport README.docx covering all six families:
  - domestic_aviation
  - international_aviation
  - domestic_shipping
  - international_shipping
  - rail_passenger
  - rail_freight

Calibrated to AES 2025 Table F1 (2023-24 column).
Output: non_road_transport_README.docx (placed alongside this script).
"""

from pathlib import Path
from datetime import date

import docx
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT_PATH = Path(__file__).parent / "non_road_transport_README.docx"


# ── Calibration data ──────────────────────────────────────────────────────────

FAMILIES = [
    dict(
        family_id="domestic_aviation",
        title="Domestic aviation",
        anchor="75,500 million pkm",
        anchor_curve="growing_then_flat__domestic_aviation",
        intensity="1,792 GJ/million_pkm (2025) → 1,592 (2050)",
        fuels="Aviation turbine fuel 100%",
        aes_total="135.3 PJ",
        aes_source="AES 2025 Table F1, sector 49 'Of which domestic air transport' (2023-24)",
        co2e_total="9.67 MtCO2e (scope 1)",
        nggi_source="NGGI 2025 Cat 1A3a domestic aviation",
    ),
    dict(
        family_id="international_aviation",
        title="International aviation",
        anchor="70,000 million pkm",
        anchor_curve="growing__international_aviation",
        intensity="2,843 GJ/million_pkm (2025) → 2,533 (2050)",
        fuels="Aviation turbine fuel 100% (bunker)",
        aes_total="199.0 PJ",
        aes_source="AES 2025 Table F1, sector 49 'Of which international air transport' (2023-24)",
        co2e_total="14.2 MtCO2e (scope 1, bunkers)",
        nggi_source="NGGI 2025 international aviation bunkers",
    ),
    dict(
        family_id="domestic_shipping",
        title="Domestic shipping (coastal)",
        anchor="30,000 million tkm",
        anchor_curve="stable__domestic_shipping",
        intensity="1,483 GJ/million_tkm (2025) → 1,320 (2050)",
        fuels="Marine diesel oil 65% / Heavy fuel oil 35%",
        aes_total="44.5 PJ",
        aes_source="AES 2025 Table F1, sector 48 'Of which coastal bunkers' (2023-24)",
        co2e_total="3.38 MtCO2e (scope 1)",
        nggi_source="NGGI 2025 domestic navigation",
    ),
    dict(
        family_id="international_shipping",
        title="International shipping",
        anchor="200,000 million tkm (estimated)",
        anchor_curve="growing__international_shipping",
        intensity="147.5 GJ/million_tkm (2025) → 125 (2050)",
        fuels="Heavy fuel oil 60% / Marine diesel oil 40% (bunker)",
        aes_total="29.5 PJ",
        aes_source="AES 2025 Table F1, sector 48 'Of which international bunkers' (2023-24)",
        co2e_total="2.3 MtCO2e (scope 1, bunkers)",
        nggi_source="NGGI 2025 international maritime bunkers",
    ),
    dict(
        family_id="rail_passenger",
        title="Rail passenger",
        anchor="22,000 million pkm",
        anchor_curve="growing__rail_passenger",
        intensity="727 GJ/million_pkm (2025) → 722 (2050)",
        fuels="Electricity 586 GJ (urban metro) / Diesel 141 GJ (regional)",
        aes_total="16.0 PJ",
        aes_source="AES 2025 Table F1, sector 47 Rail electricity 12.9 PJ + regional diesel 3.1 PJ allocated to passenger (2023-24)",
        co2e_total="0.22 MtCO2e (scope 1 diesel only)",
        nggi_source="NGGI 2025 rail (passenger share)",
    ),
    dict(
        family_id="rail_freight",
        title="Rail freight",
        anchor="700 billion tkm",
        anchor_curve="stable_growing__rail_freight",
        intensity="70,000 GJ/billion_tkm (2025) → 58,800 (2050)",
        fuels="Diesel 100% (heavy haul + interstate intermodal)",
        aes_total="49.0 PJ",
        aes_source="AES 2025 Table F1, sector 47 Rail diesel 52 PJ minus 3.1 PJ regional passenger (2023-24)",
        co2e_total="3.43 MtCO2e (scope 1)",
        nggi_source="NGGI 2025 rail (freight share)",
    ),
]

CHANGE_HISTORY = [
    ("Family", "Old AES PJ", "New AES PJ", "Source"),
    ("Domestic aviation", "85.0 (AES 2023-24)", "135.3 (AES 2025 F1)", "Revision: domestic air transport now reported separately at 135.3 PJ"),
    ("International aviation", "165.0 (AES 2023-24)", "199.0 (AES 2025 F1)", "Increase reflects post-COVID international travel rebound to 2023-24"),
    ("Domestic shipping", "30.0 (AES 2023-24)", "44.5 (AES 2025 F1)", "Coastal bunkers re-stated higher in AES 2025"),
    ("International shipping", "65.0 (AES 2023-24)", "29.5 (AES 2025 F1)", "International bunkers re-stated lower (boundary correction)"),
    ("Rail passenger", "8.0 (AES 2023-24)", "16.0 (AES 2025 F1)", "Reflects 12.9 PJ rail electricity + 3.1 PJ regional diesel"),
    ("Rail freight", "25.0 (AES 2023-24)", "49.0 (AES 2025 F1)", "Reflects 49 PJ rail diesel allocated to freight"),
]


# ── Document construction ────────────────────────────────────────────────────

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_paragraph(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p


def add_kv_table(doc, rows):
    """Two-column key/value table."""
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Light Grid Accent 1"
    for i, (k, v) in enumerate(rows):
        table.cell(i, 0).text = k
        table.cell(i, 1).text = v
        for cell in (table.cell(i, 0), table.cell(i, 1)):
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
        # Bold the key column.
        for run in table.cell(i, 0).paragraphs[0].runs:
            run.bold = True
    return table


def add_grid_table(doc, header, rows):
    """Multi-column header + rows."""
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = "Light Grid Accent 1"
    for j, label in enumerate(header):
        cell = table.cell(0, j)
        cell.text = label
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            cell = table.cell(i, j)
            cell.text = str(val)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)
    return table


def build():
    doc = Document()

    # Title
    title = doc.add_heading("Non-Road Transport Sector — Recalibration Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_paragraph(
        doc,
        f"AES 2025 Table F1 (2023-24 column) recalibration  |  {date.today().isoformat()}",
        italic=True,
        size=10,
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Overview
    add_heading(doc, "1. Overview", level=1)
    add_paragraph(
        doc,
        "This report documents the recalibration of the six non-road transport families "
        "(domestic and international aviation, domestic and international shipping, "
        "and rail passenger and rail freight) against the Australian Energy Statistics "
        "2025 (AES 2025) Table F1, using the 2023-24 column. The recalibration replaces "
        "the prior AES 2023-24 calibration that had been applied previously. "
        "Service-demand anchors (passenger-km, tonne-km) are unchanged; only the "
        "per-unit energy intensity coefficients in each family's incumbent state "
        "are updated. Transition-state intensities (SAF blends, fuel-cell aircraft, "
        "battery-electric vessels, electrified rail) remain at their technology-specific "
        "values because they reflect engineering characteristics rather than the AES baseline.",
    )

    # Calibration summary
    add_heading(doc, "2. Calibration summary (AES 2025 Table F1, 2023-24)", level=1)
    summary_header = ["Family", "Anchor", "Total energy", "Total scope 1 CO2e"]
    summary_rows = [
        [f["title"], f["anchor"], f["aes_total"], f["co2e_total"]] for f in FAMILIES
    ]
    add_grid_table(doc, summary_header, summary_rows)

    # Change history
    add_heading(doc, "3. Calibration changes vs prior AES 2023-24 calibration", level=1)
    add_grid_table(doc, list(CHANGE_HISTORY[0]), [list(r) for r in CHANGE_HISTORY[1:]])

    # Key takeaways
    add_heading(doc, "4. Key takeaways", level=2)
    bullets = [
        "Domestic aviation almost doubled (85 → 135.3 PJ); reflects the full reporting "
        "of domestic air transport in AES 2025 Table F1.",
        "International aviation bunkers rose 165 → 199 PJ; consistent with post-COVID "
        "long-haul service recovery to/from Australia.",
        "Coastal shipping rose 30 → 44.5 PJ; international bunkers fell 65 → 29.5 PJ. "
        "The international value drop is a boundary re-statement rather than a real decline.",
        "Rail total (sector 47) is now 64.9 PJ (52 PJ diesel + 12.9 PJ electricity), "
        "almost double the prior 33 PJ aggregate. The split applied here allocates "
        "all rail electricity (12.9 PJ) plus 3.1 PJ regional diesel to passenger "
        "services (16 PJ), and the remaining 49 PJ diesel to freight.",
    ]
    for b in bullets:
        p = doc.add_paragraph(b, style="List Bullet")
        for run in p.runs:
            run.font.size = Pt(11)

    # Per-family detail
    add_heading(doc, "5. Per-family calibration detail", level=1)

    for fam in FAMILIES:
        add_heading(doc, fam["title"], level=2)
        add_kv_table(
            doc,
            [
                ("Family ID", fam["family_id"]),
                ("Service demand anchor (2025)", fam["anchor"]),
                ("Demand growth curve", fam["anchor_curve"]),
                ("Energy intensity (incumbent)", fam["intensity"]),
                ("Fuel mix at 2025", fam["fuels"]),
                ("Total energy at anchor (AES 2025)", fam["aes_total"]),
                ("AES source", fam["aes_source"]),
                ("Total scope 1 CO2e at anchor", fam["co2e_total"]),
                ("NGGI cross-check", fam["nggi_source"]),
            ],
        )

    # Method notes
    add_heading(doc, "6. Method notes", level=1)
    add_paragraph(
        doc,
        "Demand anchors are unchanged because BITRE service volume statistics for "
        "passenger-km and tonne-km were not affected by the AES revision. Only "
        "per-unit energy intensity coefficients on the incumbent (conventional) "
        "state in each family were rescaled. Transition-state coefficients "
        "(SAF blends, hydrogen aircraft, battery vessels, electrified rail) preserve "
        "their technology-specific values because those reflect engineering "
        "characteristics rather than the AES baseline.",
    )
    add_paragraph(
        doc,
        "The rail passenger / freight split is not directly published in AES. "
        "The allocation applied here assigns the full 12.9 PJ rail electricity "
        "to passenger services (urban metro is electric-dominant) plus 3.1 PJ "
        "of regional diesel passenger services, totalling 16 PJ rail passenger. "
        "The remainder (49 PJ diesel) is allocated to rail freight. This carries "
        "calibration uncertainty until a Phase 2 disaggregation by service class.",
    )

    # Sources
    add_heading(doc, "7. Sources", level=1)
    add_paragraph(doc, "S001 — Australian Energy Statistics 2025 (Table F1, 2023-24 column).")
    add_paragraph(doc, "S004 — National Greenhouse Gas Inventory (NGGI) 2025.")
    add_paragraph(doc, "S012 — BITRE Aviation Statistical Report 2023-24; Rail Summary Data 2023-24.")
    add_paragraph(doc, "S013 — BITRE coastal freight statistics; freight linehaul.")
    add_paragraph(doc, "S032 — NGA Factors 2025 (NGA 2024) — emission factors for diesel, jet fuel, MDO, HFO.")

    # Files affected
    add_heading(doc, "8. Files affected", level=1)
    files_changed = [
        "sector_trajectory_library/families/domestic_aviation/generate_domestic_aviation_data.py",
        "sector_trajectory_library/families/domestic_aviation/family_states.csv (regenerated)",
        "sector_trajectory_library/families/domestic_aviation/demand.csv (regenerated)",
        "sector_trajectory_library/families/domestic_aviation/README.md",
        "sector_trajectory_library/families/international_aviation/generate_international_aviation_data.py",
        "sector_trajectory_library/families/international_aviation/family_states.csv (regenerated)",
        "sector_trajectory_library/families/international_aviation/demand.csv (regenerated)",
        "sector_trajectory_library/families/international_aviation/README.md",
        "sector_trajectory_library/families/domestic_shipping/generate_domestic_shipping_data.py",
        "sector_trajectory_library/families/domestic_shipping/family_states.csv (regenerated)",
        "sector_trajectory_library/families/domestic_shipping/demand.csv (regenerated)",
        "sector_trajectory_library/families/domestic_shipping/README.md",
        "sector_trajectory_library/families/international_shipping/generate_international_shipping_data.py",
        "sector_trajectory_library/families/international_shipping/family_states.csv (regenerated)",
        "sector_trajectory_library/families/international_shipping/demand.csv (regenerated)",
        "sector_trajectory_library/families/international_shipping/README.md",
        "sector_trajectory_library/families/rail_passenger/generate_rail_passenger_data.py",
        "sector_trajectory_library/families/rail_passenger/family_states.csv (regenerated)",
        "sector_trajectory_library/families/rail_passenger/demand.csv (regenerated)",
        "sector_trajectory_library/families/rail_passenger/README.md",
        "sector_trajectory_library/families/rail_freight/generate_rail_freight_data.py",
        "sector_trajectory_library/families/rail_freight/family_states.csv (regenerated)",
        "sector_trajectory_library/families/rail_freight/demand.csv (regenerated)",
        "sector_trajectory_library/families/rail_freight/README.md",
        "web/test/firstClassResiduals.test.mjs (electricity balance assertion updated)",
    ]
    for f in files_changed:
        p = doc.add_paragraph(f, style="List Bullet")
        for run in p.runs:
            run.font.size = Pt(10)

    # Save
    doc.save(str(OUT_PATH))
    print(f"Wrote {OUT_PATH}")


if __name__ == "__main__":
    build()
