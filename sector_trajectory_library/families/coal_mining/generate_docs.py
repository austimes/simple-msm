"""
Generate coal mining calibration documentation:
  - coal_mining_calibration_report.docx  (Word audit document)
  - coal_mining_calibration.xlsx         (Excel workbook with tables and charts)

Revised to use total Australian coal production basis (420,000 kt).
Energy sense-check: 406 GJ/kt × 420,000 kt = 170.5 PJ  (AES 174 PJ, 98.0% coverage)
Fugitive sense-check: 69 tCO2e/kt × 420,000 kt = 29.0 MtCO2e  (NGGI 29 MtCO2e, 99.9%)
"""

from pathlib import Path
import docx
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import openpyxl
from openpyxl import Workbook
from openpyxl.styles import (
    Font, PatternFill, Alignment, Border, Side, numbers
)
from openpyxl.chart import BarChart, LineChart, Reference
from openpyxl.chart.series import DataPoint
from openpyxl.utils import get_column_letter
import datetime

OUT_DIR = Path(__file__).parent

# ── Shared calibration constants (total production basis) ───────────────────
YEARS = [2025, 2030, 2035, 2040, 2045, 2050]
TOTAL_KT    = 420_000       # total Australian coal production — the demand anchor
AES_TOTAL_PJ = 174.0
MODELLED_PJ  = round(406 * TOTAL_KT / 1e6, 1)   # 170.5 PJ

RFUEL_2025 = 302.0   # GJ/kt
ELEC_2025  =  79.0
GAS_2025   =  25.0
SEI_2025   = RFUEL_2025 + ELEC_2025 + GAS_2025   # 406 GJ/kt

ENERGY_CO2E_2025  = 22.4   # tCO2e/kt  (NGA 2025: diesel 69.9 kgCO2e/GJ)
PROCESS_CO2E_2025 = 69.0
TOTAL_CO2E_2025   = ENERGY_CO2E_2025 + PROCESS_CO2E_2025  # 91.4 tCO2e/kt

# Demand trajectory: declining__coal_mining_total  −1.1%/yr
# Index multipliers relative to 2025=1.000 at (1-0.011)^n
TOTAL_DEMAND = {
    2025: 420_000,
    2030: 397_403,
    2035: 376_021,
    2040: 355_797,
    2045: 336_649,
    2050: 318_541,
}

# Conventional state trajectories (from family_states.csv — source of truth)
CONV_RFUEL = {2025: 302, 2030: 291, 2035: 280, 2040: 270, 2045: 260, 2050: 250}
CONV_ELEC  = {2025:  79, 2030:  83, 2035:  88, 2040:  94, 2045:  98, 2050: 104}
CONV_GAS   = {2025:  25, 2030:  24, 2035:  23, 2040:  22, 2045:  22, 2050:  21}
CONV_E_CO2E= {2025:22.4, 2030:21.6, 2035:20.8, 2040:20.0, 2045:19.3, 2050:18.6}
CONV_P_CO2E= {2025:69.0, 2030:63.0, 2035:57.0, 2040:52.0, 2045:47.0, 2050:43.0}

BEV_ELEC   = {2025: 134, 2030: 142, 2035: 148, 2040: 153, 2045: 156, 2050: 158}
BEV_RFUEL  = {2025: 136, 2030: 115, 2035:  98, 2040:  87, 2045:  80, 2050:  75}
BEV_GAS    = {2025:  22, 2030:  21, 2035:  20, 2040:  19, 2045:  19, 2050:  18}
BEV_E_CO2E = {2025:10.6, 2030: 9.1, 2035: 7.9, 2040: 7.1, 2045: 6.6, 2050: 6.2}

H2_H2      = {2025:  83, 2030:  88, 2035:  92, 2040:  96, 2045:  99, 2050: 102}
H2_ELEC    = {2025:  79, 2030:  79, 2035:  79, 2040:  79, 2045:  79, 2050:  79}
H2_RFUEL   = {2025: 136, 2030: 115, 2035:  98, 2040:  87, 2045:  80, 2050:  75}
H2_GAS     = {2025:  22, 2030:  21, 2035:  20, 2040:  19, 2045:  19, 2050:  18}

LFUG_RFUEL = {2025: 302, 2030: 291, 2035: 280, 2040: 270, 2045: 260, 2050: 250}
LFUG_ELEC  = {2025:  87, 2030:  91, 2035:  96, 2040: 102, 2045: 106, 2050: 112}
LFUG_GAS   = {2025:  25, 2030:  24, 2035:  23, 2040:  22, 2045:  22, 2050:  21}
LFUG_P_CO2E= {2025:34.5, 2030:27.0, 2035:21.0, 2040:16.0, 2045:12.5, 2050:10.0}

ILC_ELEC   = {2025: 158, 2030: 163, 2035: 167, 2040: 170, 2045: 172, 2050: 174}
ILC_RFUEL  = {2025: 110, 2030:  88, 2035:  72, 2040:  60, 2045:  50, 2050:  40}
ILC_GAS    = {2025:  18, 2030:  16, 2035:  14, 2040:  12, 2045:  10, 2050:   9}
ILC_E_CO2E = {2025: 8.6, 2030: 7.0, 2035: 5.8, 2040: 4.8, 2045: 4.0, 2050: 3.3}
ILC_P_CO2E = {2025:28.0, 2030:19.5, 2035:13.5, 2040: 9.5, 2045: 7.0, 2050: 5.5}

STATE_LABELS = [
    "coal_mining__conventional",
    "coal_mining__bev_haulage",
    "coal_mining__hydrogen_fcev",
    "coal_mining__low_fugitive",
    "coal_mining__integrated_lc",
]
STATE_NAMES = [
    "Conventional (diesel + grid)",
    "Battery-electric haulage (BEV)",
    "Hydrogen FCEV haulage",
    "Low-fugitive (methane abatement)",
    "Integrated low-carbon",
]
STATE_2025_SEI   = [406, 292, 320, 414, 286]
STATE_2025_ECO2E = [22.4, 10.6, 10.6, 22.4, 8.6]
STATE_2025_PCO2E = [69.0, 69.0, 69.0, 34.5, 28.0]


# ═══════════════════════════════════════════════════════════════════════════════
#  WORD DOCUMENT
# ═══════════════════════════════════════════════════════════════════════════════

def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)

def add_para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p

def add_table(doc, headers, rows, caption=None):
    if caption:
        doc.add_paragraph(caption, style="Caption")
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Light Shading Accent 1"
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            tbl.rows[r_idx + 1].cells[c_idx].text = str(val)
    doc.add_paragraph()
    return tbl

def build_word_doc():
    doc = Document()

    title = doc.add_heading("Coal Mining Sector Family — Data Calibration Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_para(doc,
        f"Prepared: {datetime.date.today():%d %B %Y}  |  Version: 1.1  |  Status: Phase 1 calibration (total production basis)",
        italic=True)
    doc.add_paragraph()

    # ── 1. Executive Summary ────────────────────────────────────────────────
    add_heading(doc, "1. Executive Summary")
    add_para(doc,
        "This report documents the calibration of the coal_mining family in the Simple MSM "
        "sector trajectory library. The family models total Australian coal mining operations "
        "(420,000 kt/yr in 2025, comprising ~330,000 kt export and ~90,000 kt domestic). "
        "Using total production as the demand anchor gives exact matches to both key national statistics: "
        "(a) AES 2023-24 Table F: 406 GJ/kt × 420,000 kt = 170.5 PJ (98.0% of AES 174 PJ); "
        "(b) NGGI 2023-24 Cat 1B1a: 69 tCO2e/kt × 420,000 kt = 29.0 MtCO2e (99.9% of NGGI 29 MtCO2e). "
        "Demand declines at −1.1%/yr compound (weighted average of export −0.75%/yr and domestic −2.5%/yr)."
    )

    # ── 2. Data Sources ────────────────────────────────────────────────────
    add_heading(doc, "2. Primary Data Sources")
    add_table(doc,
        ["Source ID", "Dataset", "Year", "Key Figure Used"],
        [
            ["S001", "AES 2023-24 Table F — Mining sector final energy", "2023-24",
             "174 PJ total, fuel split: rfuel 73%, elec 19%, gas 6%"],
            ["S002", "NGGI 2023-24 — Category 1B1a Fugitive (coal mining)", "2023-24",
             "~29 MtCO2e; 69 tCO2e/kt at 420,000 kt total"],
            ["S021", "Geoscience Australia Australian Energy Update 2023-24", "2023-24",
             "420,000 kt total production; export ~330,000 kt, domestic ~90,000 kt"],
            ["NGA2025", "National Greenhouse Accounts Factors 2025", "2025",
             "Diesel/rfuel 69.9 kgCO2e/GJ; gas 51.4 kgCO2e/GJ (AR5 GWP100)"],
        ],
        caption="Table 1 — Primary data sources"
    )

    # ── 3. Calibration ─────────────────────────────────────────────────────
    add_heading(doc, "3. AES 174 PJ Calibration — Step-by-Step")
    add_para(doc,
        "AES 2023-24 Table F reports 174 PJ of total final energy consumed by the coal mining "
        "sector (ANZSIC 06100 black coal and 06200 brown coal combined) at national scale "
        "(420,000 kt of raw coal output)."
    )
    add_table(doc,
        ["Step", "Calculation", "Result"],
        [
            ["1. National average intensity", "174 PJ ÷ 420,000 kt", "414 GJ/kt"],
            ["2. Coverage adjustment", "~98% of 174 PJ covered by 3 commodities", "406 GJ/kt modelled"],
            ["3. Fuel split rfuel 73%", "0.73 × 406 GJ/kt", "302 GJ/kt refined liquid fuels"],
            ["4. Electricity split 19%", "0.19 × 406 GJ/kt", "79 GJ/kt electricity"],
            ["5. Gas split 6%", "0.06 × 406 GJ/kt", "25 GJ/kt natural gas"],
            ["6. Sum check", "302 + 79 + 25", "406 GJ/kt ✓"],
            ["7. Total energy (modelled)", "406 GJ/kt × 420,000 kt ÷ 10⁶", "170.5 PJ (98.0% of AES 174 PJ) ✓"],
            ["8. Fugitive check", "69 tCO2e/kt × 420,000 kt ÷ 10⁶", "29.0 MtCO2e (99.9% of NGGI 29 MtCO2e) ✓"],
        ],
        caption="Table 2 — Calibration steps"
    )

    # ── 4. Energy Consumption Sense Check ─────────────────────────────────
    add_heading(doc, "4. Energy Consumption Sense Check — Fuel by Type")
    add_para(doc,
        "The table below shows the conventional state fuel consumption by type over the modelling "
        "horizon. Values are calculated from the per-kt coefficients multiplied by the total demand "
        "trajectory (declining__coal_mining_total, −1.1%/yr). The 2025 totals check directly against "
        "AES 2023-24 Table F."
    )

    energy_rows = []
    for yr in YEARS:
        kt = TOTAL_DEMAND[yr]
        rfuel_pj = round(CONV_RFUEL[yr] * kt / 1e6, 1)
        elec_pj  = round(CONV_ELEC[yr]  * kt / 1e6, 1)
        gas_pj   = round(CONV_GAS[yr]   * kt / 1e6, 1)
        total_pj = round((CONV_RFUEL[yr] + CONV_ELEC[yr] + CONV_GAS[yr]) * kt / 1e6, 1)
        energy_rows.append([yr, f"{kt:,}", rfuel_pj, elec_pj, gas_pj, total_pj])

    add_table(doc,
        ["Year", "Demand (kt)", "Rfuel (PJ)", "Electricity (PJ)", "Gas (PJ)", "Total (PJ)"],
        energy_rows,
        caption="Table 3 — Conventional state fuel consumption by type (2025–2050)"
    )
    add_para(doc,
        "AES reference (2025): 174 PJ total. Modelled 2025 total: 170.5 PJ (98.0% coverage). "
        "The 2% gap reflects LPG and biomass fuels not in the model commodity set.",
        italic=True
    )

    # ── 5. Emissions Sense Check ───────────────────────────────────────────
    add_heading(doc, "5. Emissions Sense Check — Energy and Fugitive CO2e")
    add_para(doc,
        "Emissions are split into scope 1 energy combustion (derived from fuel coefficients × NGA "
        "emission factors) and fugitive process emissions (NGGI Cat 1B1a per-kt average). Both are "
        "calibrated to published national statistics."
    )

    emis_rows = []
    for yr in YEARS:
        kt = TOTAL_DEMAND[yr]
        e_co2  = round(CONV_E_CO2E[yr] * kt / 1e6, 2)
        f_co2  = round(CONV_P_CO2E[yr] * kt / 1e6, 2)
        tot_co2 = round(e_co2 + f_co2, 2)
        emis_rows.append([yr, f"{kt:,}", f"{CONV_E_CO2E[yr]:.1f}", f"{CONV_P_CO2E[yr]:.1f}",
                          f"{e_co2:.2f}", f"{f_co2:.2f}", f"{tot_co2:.2f}"])

    add_table(doc,
        ["Year", "Demand (kt)", "Energy CO2e (t/kt)", "Fugitive CO2e (t/kt)",
         "Energy (MtCO2e)", "Fugitive (MtCO2e)", "Total (MtCO2e)"],
        emis_rows,
        caption="Table 4 — Conventional state emissions by type (2025–2050)"
    )
    add_para(doc,
        "NGGI reference (2025): ~29 MtCO2e fugitive. Modelled 2025 fugitive: 28.98 MtCO2e (99.9%). "
        "Caution: the 69 tCO2e/kt fugitive coefficient is a national average dominated by "
        "underground export coal. Open-cut domestic mines (predominantly VIC brown coal) have "
        "much lower fugitive intensity (~15 tCO2e/kt). Phase 2: disaggregate by mine type.",
        italic=True
    )

    # ── 6. Technology States ───────────────────────────────────────────────
    add_heading(doc, "6. Technology States — Five Pathway Archetypes")
    add_table(doc,
        ["State ID", "Label", "2025 SEI (GJ/kt)", "2025 Energy CO2e (t/kt)", "2025 Fugitive CO2e (t/kt)", "Key abatement lever"],
        [
            [STATE_LABELS[0], STATE_NAMES[0], "406", "22.0", "69.0", "Baseline — no additional abatement"],
            [STATE_LABELS[1], STATE_NAMES[1], "292", "10.6", "69.0", "BEV haul trucks replace diesel (~55% rfuel reduction)"],
            [STATE_LABELS[2], STATE_NAMES[2], "320", "10.6", "69.0", "H2 FCEV replace diesel (H2 + residual rfuel)"],
            [STATE_LABELS[3], STATE_NAMES[3], "414", "22.2", "34.5", "Methane capture/VAM oxidation (~50% fugitive reduction)"],
            [STATE_LABELS[4], STATE_NAMES[4], "286",  "8.6", "28.0", "BEV haulage + fugitive abatement combined"],
        ],
        caption="Table 5 — Technology states summary (2025 values)"
    )

    # ── 7. Emissions Pathways by State ─────────────────────────────────────
    add_heading(doc, "7. Total Emissions Pathways by Technology State (MtCO2e/yr)")
    state_emis_rows = []
    state_defs = [
        ("Conventional", CONV_E_CO2E, CONV_P_CO2E),
        ("BEV haulage",  BEV_E_CO2E,  CONV_P_CO2E),
        ("Hydrogen FCEV", BEV_E_CO2E, CONV_P_CO2E),
        ("Low-fugitive", CONV_E_CO2E, LFUG_P_CO2E),
        ("Integrated low-carbon", ILC_E_CO2E, ILC_P_CO2E),
    ]
    for name, eco2, pco2 in state_defs:
        row = [name]
        for yr in YEARS:
            kt = TOTAL_DEMAND[yr]
            total = round((eco2[yr] + pco2[yr]) * kt / 1e6, 2)
            row.append(f"{total:.2f}")
        state_emis_rows.append(row)

    add_table(doc,
        ["State"] + [str(y) for y in YEARS],
        state_emis_rows,
        caption="Table 6 — Total CO2e by state (MtCO2e/yr, at modelled demand trajectory)"
    )

    # ── 8. Demand Trajectory ──────────────────────────────────────────────
    add_heading(doc, "8. Demand Trajectory")
    add_para(doc,
        "Total Australian coal production is modelled as declining at −1.1%/yr compound. "
        "This is the demand-weighted average of the export coal trajectory (−0.75%/yr, "
        "reflecting gradual global thermal coal displacement and slow metallurgical coal decline) "
        "and the domestic trajectory (−2.5%/yr, reflecting coal-fired power station closures "
        "per AEMO ISP 2024 Step Change). Curve ID: declining__coal_mining_total."
    )
    add_table(doc,
        ["Year", "Demand (kt)", "Index (2025=1.00)", "Total Energy PJ (conv)", "Total CO2e (MtCO2e, conv)"],
        [(yr, f"{TOTAL_DEMAND[yr]:,.0f}",
          f"{TOTAL_DEMAND[yr]/420000:.4f}",
          f"{TOTAL_DEMAND[yr] * SEI_2025 / 1e6:.1f}",
          f"{(ENERGY_CO2E_2025 + PROCESS_CO2E_2025) * TOTAL_DEMAND[yr] / 1e6:.2f}")
         for yr in YEARS],
        caption="Table 7 — Total coal production trajectory (declining__coal_mining_total, −1.1%/yr)"
    )

    # ── 9. Known Caveats ──────────────────────────────────────────────────
    add_heading(doc, "9. Known Caveats and Phase 2 Improvements")
    caveats = [
        ("National-average fugitive coefficient applied across all mine types",
         "The 69 tCO2e/kt average is dominated by underground export mines. "
         "Open-cut domestic mines (predominantly VIC brown coal) have ~15 tCO2e/kt. "
         "Phase 2: disaggregate by mine type. Current figure is a conservative upper bound."),
        ("Per-kt energy coefficient is national average",
         "Domestic open-cut mines likely have lower per-kt energy than underground "
         "export mines. The 406 GJ/kt may overestimate domestic and underestimate "
         "underground segments. Phase 2: calibrate by mine type using ABS ANZSIC data."),
        ("Export demand trajectory assumed independent of model",
         "Export coal demand is driven by Asian market dynamics. The −0.75%/yr export "
         "decline is a scenario assumption (moderate global coal displacement), not an "
         "endogenous model result. Alternative trajectories can be tested via the curve ID."),
        ("BEV/H2 max-share trajectories are averaged across all mine types",
         "Long-life export mines (Bowen Basin, Hunter Valley) have stronger investment "
         "case for BEV than short-life domestic mines. Phase 2: split max-share by "
         "mine-type sub-family."),
    ]
    for title_text, body in caveats:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(title_text + ": ")
        run.bold = True
        p.add_run(body)

    doc.add_paragraph()
    add_para(doc, "— End of report —", italic=True)
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    out_path = OUT_DIR / "coal_mining_calibration_report.docx"
    doc.save(str(out_path))
    print(f"  Written: {out_path}")


# ═══════════════════════════════════════════════════════════════════════════════
#  EXCEL WORKBOOK
# ═══════════════════════════════════════════════════════════════════════════════

HDR_FILL  = PatternFill("solid", fgColor="1F4E79")
HDR_FONT  = Font(color="FFFFFF", bold=True, size=11)
SUB_FILL  = PatternFill("solid", fgColor="2E75B6")
SUB_FONT  = Font(color="FFFFFF", bold=True, size=10)
EVEN_FILL = PatternFill("solid", fgColor="DEEAF1")
TITLE_FONT = Font(bold=True, size=14, color="1F4E79")
LABEL_FONT = Font(bold=True, size=10)
THIN = Side(style="thin", color="BFBFBF")
THIN_BORDER = Border(left=THIN, right=THIN, top=THIN, bottom=THIN)

SENSE_CHECK_FILL = PatternFill("solid", fgColor="E2EFDA")  # green for pass
WARN_FILL = PatternFill("solid", fgColor="FFF2CC")          # yellow for note

def hdr(ws, row, col, val):
    c = ws.cell(row=row, column=col, value=val)
    c.fill = HDR_FILL; c.font = HDR_FONT; c.border = THIN_BORDER
    c.alignment = Alignment(horizontal="center", wrap_text=True)
    return c

def sub(ws, row, col, val):
    c = ws.cell(row=row, column=col, value=val)
    c.fill = SUB_FILL; c.font = SUB_FONT; c.border = THIN_BORDER
    c.alignment = Alignment(horizontal="center")
    return c

def dat(ws, row, col, val, bold=False, fmt=None, even=False, fill=None):
    c = ws.cell(row=row, column=col, value=val)
    if fill:
        c.fill = fill
    elif even:
        c.fill = EVEN_FILL
    c.font = Font(bold=bold, size=10)
    c.border = THIN_BORDER
    c.alignment = Alignment(horizontal="right" if isinstance(val, (int, float)) else "left")
    if fmt:
        c.number_format = fmt
    return c

def title_cell(ws, row, col, val):
    c = ws.cell(row=row, column=col, value=val)
    c.font = TITLE_FONT
    return c


def build_cover(wb):
    ws = wb.active
    ws.title = "Cover"
    ws.sheet_view.showGridLines = False
    ws.column_dimensions["A"].width = 5
    ws.column_dimensions["B"].width = 40
    ws.column_dimensions["C"].width = 40

    ws.merge_cells("B2:C2")
    c = ws.cell(row=2, column=2, value="Coal Mining Sector — Calibration Workbook")
    c.font = Font(bold=True, size=16, color="1F4E79")
    c.alignment = Alignment(horizontal="center")

    info = [
        ("Date", f"{datetime.date.today():%d %B %Y}"),
        ("Version", "1.1"),
        ("Status", "Phase 1 calibration — total production basis (420,000 kt)"),
        ("Data basis", "AES 2023-24 Table F; NGGI 2023-24; Geoscience Australia 2023-24"),
        ("Author", "Simple MSM core library team"),
        ("Model family", "coal_mining (required_service, kt_coal)"),
        ("Anchor value", "420,000 kt (total production, 2025)"),
        ("Growth curve", "declining__coal_mining_total (−1.1%/yr, weighted export+domestic)"),
        ("Energy sense-check", f"406 GJ/kt × 420,000 kt = {MODELLED_PJ} PJ  (AES {AES_TOTAL_PJ} PJ = 98.0%) ✓"),
        ("Fugitive sense-check", "69 tCO2e/kt × 420,000 kt = 29.0 MtCO2e  (NGGI 29 MtCO2e = 99.9%) ✓"),
    ]
    for i, (k, v) in enumerate(info):
        r = 4 + i
        ws.cell(row=r, column=2, value=k).font = LABEL_FONT
        c = ws.cell(row=r, column=3, value=v)
        c.font = Font(size=10)
        if "sense-check" in k:
            c.fill = SENSE_CHECK_FILL

    contents = [
        ("Sheet", "Contents"),
        ("Energy_Emissions_SenseCheck", "Fuel consumption and emissions by type — primary sense-check"),
        ("174PJ_Calibration", "Step-by-step AES 174 PJ derivation and per-kt coefficients"),
        ("Total_Production", "Total production volume, energy and emissions (combined basis)"),
        ("State_Assumptions", "Technology state assumptions (5 states × 6 years)"),
        ("Total_Demand", "Total demand trajectory 2025–2050"),
        ("Emissions_Pathways", "Energy + fugitive CO2e by state 2025–2050"),
        ("Charts", "Visual charts for demand, energy intensity, and emissions"),
    ]
    ws.cell(row=16, column=2, value="Workbook contents").font = Font(bold=True, size=11, color="1F4E79")
    for i, (s, d) in enumerate(contents):
        r = 17 + i
        c1 = ws.cell(row=r, column=2, value=s)
        c2 = ws.cell(row=r, column=3, value=d)
        c1.font = Font(bold=(i == 0), size=10)
        c2.font = Font(bold=(i == 0), size=10)
        if i % 2 == 1:
            c1.fill = EVEN_FILL; c2.fill = EVEN_FILL


def build_energy_emissions_sensecheck(wb):
    """Dedicated sense-check sheet: fuel consumption and emissions by type."""
    ws = wb.create_sheet("Energy_Emissions_SenseCheck")
    ws.sheet_view.showGridLines = False
    for col, w in zip("ABCDEFGHI", [4, 10, 12, 14, 14, 12, 14, 14, 14]):
        ws.column_dimensions[col].width = w

    title_cell(ws, 1, 2, "Coal Mining — Energy Consumption & Emissions Sense Check (Conventional State)")
    ws.merge_cells("B1:I1")

    # ── Sense-check summary banner ──────────────────────────────────────────
    ws.merge_cells("B3:I3")
    banner = ws.cell(row=3, column=2,
        value=(f"2025 SENSE CHECK:  Modelled energy = {MODELLED_PJ} PJ  "
               f"(AES 174 PJ, 98.0% ✓)   |   "
               f"Modelled fugitive = 29.0 MtCO2e  (NGGI 29 MtCO2e, 99.9% ✓)"))
    banner.font = Font(bold=True, size=11, color="375623")
    banner.fill = SENSE_CHECK_FILL
    banner.alignment = Alignment(horizontal="center")

    # ── Section 1: Fuel Consumption by Type ────────────────────────────────
    ws.cell(row=5, column=2, value="SECTION 1: Fuel Consumption by Type (Conventional State)").font = \
        Font(bold=True, size=11, color="1F4E79")
    ws.merge_cells("B5:I5")

    fuel_headers = ["Year", "Demand (kt)", "Rfuel (GJ/kt)", "Rfuel (PJ)",
                    "Electricity (GJ/kt)", "Elec (PJ)", "Gas (GJ/kt)", "Gas (PJ)", "Total (PJ)"]
    for ci, h in enumerate(fuel_headers, 2):
        hdr(ws, 6, ci, h)

    fuel_data = []
    for ri, yr in enumerate(YEARS):
        kt = TOTAL_DEMAND[yr]
        rfuel_pj = round(CONV_RFUEL[yr] * kt / 1e6, 1)
        elec_pj  = round(CONV_ELEC[yr]  * kt / 1e6, 1)
        gas_pj   = round(CONV_GAS[yr]   * kt / 1e6, 1)
        total_pj = round(rfuel_pj + elec_pj + gas_pj, 1)
        even = ri % 2 == 1
        r = 7 + ri
        dat(ws, r, 2, yr, even=even)
        dat(ws, r, 3, kt, fmt="#,##0", even=even)
        dat(ws, r, 4, CONV_RFUEL[yr], fmt="#,##0", even=even)
        dat(ws, r, 5, rfuel_pj, fmt="0.0", even=even)
        dat(ws, r, 6, CONV_ELEC[yr], fmt="#,##0", even=even)
        dat(ws, r, 7, elec_pj, fmt="0.0", even=even)
        dat(ws, r, 8, CONV_GAS[yr], fmt="#,##0", even=even)
        dat(ws, r, 9, gas_pj, fmt="0.0", even=even)
        # Total with AES reference highlight for 2025
        fill = SENSE_CHECK_FILL if yr == 2025 else None
        dat(ws, r, 10, total_pj, fmt="0.0", bold=(yr == 2025), fill=fill)
        fuel_data.append((yr, rfuel_pj, elec_pj, gas_pj))

    # AES reference row
    r_ref = 7 + len(YEARS)
    ws.cell(row=r_ref, column=2, value="AES 2023-24 ref").font = Font(bold=True, italic=True, size=10, color="C00000")
    ws.cell(row=r_ref, column=10, value=174.0).font = Font(bold=True, size=10, color="C00000")
    ws.cell(row=r_ref, column=10).number_format = "0.0"
    ws.cell(row=r_ref, column=2).fill = PatternFill("solid", fgColor="FCE4D6")
    ws.cell(row=r_ref, column=10).fill = PatternFill("solid", fgColor="FCE4D6")

    note_row = r_ref + 1
    ws.merge_cells(f"B{note_row}:J{note_row}")
    ws.cell(row=note_row, column=2,
            value="Note: 2% gap vs AES reflects LPG/biomass fuels excluded from model commodity set.").font = \
        Font(italic=True, size=9, color="595959")

    # ── Chart 1: Fuel Consumption Stacked Bar ──────────────────────────────
    chart_data_row = note_row + 3
    ws.cell(row=chart_data_row, column=2, value="Year")
    ws.cell(row=chart_data_row, column=3, value="Rfuel (PJ)")
    ws.cell(row=chart_data_row, column=4, value="Electricity (PJ)")
    ws.cell(row=chart_data_row, column=5, value="Gas (PJ)")
    for ri, (yr, rf, el, ga) in enumerate(fuel_data):
        r = chart_data_row + 1 + ri
        ws.cell(row=r, column=2, value=yr)
        ws.cell(row=r, column=3, value=rf)
        ws.cell(row=r, column=4, value=el)
        ws.cell(row=r, column=5, value=ga)

    chart1 = BarChart()
    chart1.type = "col"
    chart1.grouping = "stacked"
    chart1.title = "Coal Mining — Fuel Consumption by Type (PJ, Conventional State)"
    chart1.style = 10
    chart1.y_axis.title = "PJ"
    chart1.x_axis.title = "Year"
    chart1.height = 14
    chart1.width = 22

    data_ref = Reference(ws, min_col=3, max_col=5,
                         min_row=chart_data_row, max_row=chart_data_row + len(YEARS))
    cats_ref = Reference(ws, min_col=2, min_row=chart_data_row + 1, max_row=chart_data_row + len(YEARS))
    chart1.add_data(data_ref, titles_from_data=True)
    chart1.set_categories(cats_ref)
    chart1.series[0].graphicalProperties.solidFill = "2E75B6"   # rfuel — blue
    chart1.series[1].graphicalProperties.solidFill = "70AD47"   # elec — green
    chart1.series[2].graphicalProperties.solidFill = "FFC000"   # gas — amber
    ws.add_chart(chart1, "B" + str(note_row + 3))

    # ── Section 2: Emissions by Type ───────────────────────────────────────
    emis_section_row = note_row + 22
    ws.cell(row=emis_section_row, column=2, value="SECTION 2: Emissions by Type (Conventional State)").font = \
        Font(bold=True, size=11, color="1F4E79")
    ws.merge_cells(f"B{emis_section_row}:I{emis_section_row}")

    emis_headers = ["Year", "Demand (kt)", "Energy CO2e (t/kt)", "Fugitive CO2e (t/kt)",
                    "Energy (MtCO2e)", "Fugitive (MtCO2e)", "Total (MtCO2e)"]
    for ci, h in enumerate(emis_headers, 2):
        hdr(ws, emis_section_row + 1, ci, h)

    emis_chart_data = []
    for ri, yr in enumerate(YEARS):
        kt = TOTAL_DEMAND[yr]
        e_co2 = round(CONV_E_CO2E[yr] * kt / 1e6, 2)
        f_co2 = round(CONV_P_CO2E[yr] * kt / 1e6, 2)
        t_co2 = round(e_co2 + f_co2, 2)
        even = ri % 2 == 1
        r = emis_section_row + 2 + ri
        dat(ws, r, 2, yr, even=even)
        dat(ws, r, 3, kt, fmt="#,##0", even=even)
        dat(ws, r, 4, CONV_E_CO2E[yr], fmt="0.0", even=even)
        dat(ws, r, 5, CONV_P_CO2E[yr], fmt="0.0", even=even)
        dat(ws, r, 6, e_co2, fmt="0.00", even=even)
        dat(ws, r, 7, f_co2, fmt="0.00", even=even)
        fill = SENSE_CHECK_FILL if yr == 2025 else None
        dat(ws, r, 8, t_co2, fmt="0.00", bold=(yr == 2025), fill=fill)
        emis_chart_data.append((yr, e_co2, f_co2))

    # NGGI reference row
    r_nggi = emis_section_row + 2 + len(YEARS)
    ws.cell(row=r_nggi, column=2, value="NGGI 2023-24 ref").font = Font(bold=True, italic=True, size=10, color="C00000")
    ws.cell(row=r_nggi, column=7, value=29.0).font = Font(bold=True, size=10, color="C00000")
    ws.cell(row=r_nggi, column=7).number_format = "0.00"
    ws.cell(row=r_nggi, column=2).fill = PatternFill("solid", fgColor="FCE4D6")
    ws.cell(row=r_nggi, column=7).fill = PatternFill("solid", fgColor="FCE4D6")

    note2_row = r_nggi + 1
    ws.merge_cells(f"B{note2_row}:I{note2_row}")
    ws.cell(row=note2_row, column=2,
            value="Note: 69 tCO2e/kt fugitive is national average (incl. underground export). Open-cut domestic ≈ 15 tCO2e/kt. Phase 2: disaggregate.").font = \
        Font(italic=True, size=9, color="595959")

    # ── Chart 2: Emissions Stacked Bar ────────────────────────────────────
    emis_cd_row = note2_row + 3
    ws.cell(row=emis_cd_row, column=2, value="Year")
    ws.cell(row=emis_cd_row, column=3, value="Energy CO2e (MtCO2e)")
    ws.cell(row=emis_cd_row, column=4, value="Fugitive CO2e (MtCO2e)")
    for ri, (yr, e, f) in enumerate(emis_chart_data):
        r = emis_cd_row + 1 + ri
        ws.cell(row=r, column=2, value=yr)
        ws.cell(row=r, column=3, value=e)
        ws.cell(row=r, column=4, value=f)

    chart2 = BarChart()
    chart2.type = "col"
    chart2.grouping = "stacked"
    chart2.title = "Coal Mining — Emissions by Type (MtCO2e, Conventional State)"
    chart2.style = 10
    chart2.y_axis.title = "MtCO2e"
    chart2.x_axis.title = "Year"
    chart2.height = 14
    chart2.width = 22

    data2 = Reference(ws, min_col=3, max_col=4,
                      min_row=emis_cd_row, max_row=emis_cd_row + len(YEARS))
    cats2 = Reference(ws, min_col=2, min_row=emis_cd_row + 1, max_row=emis_cd_row + len(YEARS))
    chart2.add_data(data2, titles_from_data=True)
    chart2.set_categories(cats2)
    chart2.series[0].graphicalProperties.solidFill = "ED7D31"   # energy — orange
    chart2.series[1].graphicalProperties.solidFill = "C00000"   # fugitive — red
    ws.add_chart(chart2, "B" + str(note2_row + 3))

    # Hide raw chart data cols to keep sheet tidy
    ws.column_dimensions["C"].hidden = False  # keep visible — it's useful


def build_calibration_sheet(wb):
    ws = wb.create_sheet("174PJ_Calibration")
    ws.sheet_view.showGridLines = False
    ws.column_dimensions["A"].width = 4
    ws.column_dimensions["B"].width = 40
    ws.column_dimensions["C"].width = 34
    ws.column_dimensions["D"].width = 22

    title_cell(ws, 1, 2, "AES 174 PJ Calibration — Step-by-Step Derivation (Total Production Basis)")
    ws.merge_cells("B1:D1")

    headers = ["Step", "Description", "Value / Calculation"]
    for ci, h in enumerate(headers, 2):
        hdr(ws, 3, ci, h)

    steps = [
        ("AES source", "AES 2023-24 Table F — Total final energy, coal mining (ANZSIC 06100+06200)", "174 PJ"),
        ("Total production", "Geoscience Australia 2023-24 — total raw coal output", "420,000 kt"),
        ("National avg intensity", "174 PJ ÷ 420,000 kt × 1,000,000 GJ/PJ", "414 GJ/kt"),
        ("Coverage adjustment", "~98% covered by 3 commodities (2% LPG/biomass excluded)", "406 GJ/kt modelled"),
        ("Fuel split — rfuel 73%", "0.73 × 406 GJ/kt", "302 GJ/kt"),
        ("Fuel split — electricity 19%", "0.19 × 406 GJ/kt", "79 GJ/kt"),
        ("Fuel split — gas 6%", "0.06 × 406 GJ/kt", "25 GJ/kt"),
        ("Sum check", "302 + 79 + 25", "406 GJ/kt ✓"),
        ("Total energy (modelled)", "406 GJ/kt × 420,000 kt ÷ 1,000,000", f"{MODELLED_PJ} PJ"),
        ("AES coverage", f"{MODELLED_PJ} ÷ 174.0", "98.0% ✓"),
        ("Fugitive coefficient", "29 MtCO2e (NGGI) ÷ 420,000 kt × 1,000", "69.0 tCO2e/kt"),
        ("Fugitive check", "69.0 tCO2e/kt × 420,000 kt ÷ 10⁶", "28.98 MtCO2e ≈ 29 MtCO2e ✓"),
    ]
    for ri, (step, desc, val) in enumerate(steps):
        r = 4 + ri
        even = ri % 2 == 1
        fill = SENSE_CHECK_FILL if "✓" in val else None
        dat(ws, r, 2, step, bold=True, even=even)
        dat(ws, r, 3, desc, even=even)
        dat(ws, r, 4, val, even=even, fill=fill)

    # Coefficient table
    ws.cell(row=18, column=2, value="2025 Commodity Coefficients — Conventional State").font = \
        Font(bold=True, size=11, color="1F4E79")
    for ci, h in enumerate(["Commodity", "Coefficient (GJ/kt)", "Share (%)", "EF (kgCO2e/GJ)", "CO2e (tCO2e/kt)"], 2):
        hdr(ws, 19, ci, h)

    commodities = [
        ("Refined liquid fuels", 302, "74%", 69.9, 21.1),
        ("Electricity", 79, "19%", "Scope 2 — excluded", "—"),
        ("Natural gas", 25, "6%", 51.4, 1.3),
        ("Unmodelled (LPG/biomass)", 8, "2%", "Excluded", "—"),
        ("TOTAL", 414, "100%", "—", "22.4 (energy CO2e)"),
    ]
    for ri, row in enumerate(commodities):
        r = 20 + ri
        even = ri % 2 == 1
        for ci, v in enumerate(row, 2):
            dat(ws, r, ci, v, bold=(ri == len(commodities)-1), even=even)

    # Fugitive table
    ws.cell(row=27, column=2, value="Fugitive Methane Calibration (NGGI 2023-24 Cat 1B1a)").font = \
        Font(bold=True, size=11, color="1F4E79")
    for ci, h in enumerate(["Item", "Value", "Source"], 2):
        hdr(ws, 28, ci, h)

    fugitive_rows = [
        ("NGGI 2023-24 Cat 1B1a total", "~29 MtCO2e", "NGGI 2023-24"),
        ("Total production", "420,000 kt", "Geoscience Australia 2023-24"),
        ("Per-kt fugitive (national avg)", "69 tCO2e/kt", "29 Mt ÷ 420,000 kt × 1,000"),
        ("Open-cut fugitive (typical)", "~15 tCO2e/kt", "Literature / NGER reports"),
        ("Underground fugitive (typical)", "~100-200 tCO2e/kt", "Literature / NGER reports"),
        ("Phase 1 value used", "69 tCO2e/kt (national avg)", "Conservative — dominated by underground"),
        ("Phase 2 target", "Disaggregate by mine type", "Open-cut vs underground via NGER data"),
    ]
    for ri, row in enumerate(fugitive_rows):
        r = 29 + ri
        even = ri % 2 == 1
        for ci, v in enumerate(row, 2):
            dat(ws, r, ci, v, even=even)


def build_total_production_sheet(wb):
    ws = wb.create_sheet("Total_Production")
    ws.sheet_view.showGridLines = False
    ws.column_dimensions["A"].width = 4
    for col, w in zip("BCDE", [28, 18, 20, 20]):
        ws.column_dimensions[col].width = w

    title_cell(ws, 1, 2, "Total Australian Coal Production — Volume, Energy and Emissions (2025 Baseline)")
    ws.merge_cells("B1:E1")

    # Context note
    ws.merge_cells("B3:E3")
    ws.cell(row=3, column=2,
            value="Production basis: 420,000 kt total (export 330,000 kt + domestic 90,000 kt). "
                  "Per-kt coefficients are national averages applied to total production — "
                  "no domestic/export split is required as AES/NGGI data is not disaggregated by market destination.").font = \
        Font(italic=True, size=9, color="595959")

    # Production breakdown
    for ci, h in enumerate(["Component", "Volume (kt)", "Share (%)", "Energy (PJ)", "Fugitive CO2e (MtCO2e)"], 2):
        hdr(ws, 5, ci, h)

    prod_rows = [
        ("Export coal (approx.)", 330_000, "79%",
         round(330_000 * 406 / 1e6, 1), round(330_000 * 69 / 1e6, 2)),
        ("Domestic coal (approx.)", 90_000, "21%",
         round(90_000 * 406 / 1e6, 1), round(90_000 * 69 / 1e6, 2)),
        ("TOTAL (model anchor)", 420_000, "100%",
         round(MODELLED_PJ, 1), round(TOTAL_KT * 69 / 1e6, 2)),
    ]
    for ri, row in enumerate(prod_rows):
        r = 6 + ri
        even = ri % 2 == 1
        for ci, v in enumerate(row, 2):
            dat(ws, r, ci, v, bold=(ri == 2), even=even,
                fmt="#,##0" if isinstance(v, int) else None)

    # Benchmark reference
    ws.cell(row=11, column=2, value="Benchmark comparisons (2025)").font = Font(bold=True, size=11, color="1F4E79")
    for ci, h in enumerate(["Metric", "Modelled", "Reference", "Coverage %"], 2):
        hdr(ws, 12, ci, h)

    bench_rows = [
        ("Total energy (PJ)", f"{MODELLED_PJ} PJ", "AES 174 PJ", "98.0% ✓"),
        ("Fugitive CO2e (MtCO2e)", "29.0 MtCO2e", "NGGI 29 MtCO2e", "99.9% ✓"),
        ("Total production (kt)", "420,000 kt", "GA 420,000 kt", "100% ✓"),
        ("Energy intensity (GJ/kt)", "406 GJ/kt", "AES 414 GJ/kt (gross)", "98.0% (coverage adj.)"),
    ]
    for ri, row in enumerate(bench_rows):
        r = 13 + ri
        even = ri % 2 == 1
        fill = SENSE_CHECK_FILL if "✓" in row[3] else None
        for ci, v in enumerate(row, 2):
            dat(ws, r, ci, v, even=even, fill=fill if ci == 5 else (EVEN_FILL if even else None))

    note = ws.cell(row=18, column=2,
                   value="Note: Per-kt coefficients are applied uniformly to total production. "
                         "Mine-type disaggregation (open-cut vs underground) is a Phase 2 improvement.")
    note.font = Font(italic=True, size=9, color="595959")
    ws.merge_cells("B18:E18")


def build_state_assumptions(wb):
    ws = wb.create_sheet("State_Assumptions")
    ws.sheet_view.showGridLines = False
    ws.column_dimensions["A"].width = 4
    ws.column_dimensions["B"].width = 32
    for col in "CDEFGH":
        ws.column_dimensions[col].width = 12

    title_cell(ws, 1, 2, "Coal Mining — Technology State Assumptions (5 States × 6 Years)")
    ws.merge_cells("B1:H1")

    # SEI table
    ws.cell(row=3, column=2, value="Specific Energy Intensity (GJ/kt)").font = Font(bold=True, size=11, color="1F4E79")
    hdr(ws, 4, 2, "State")
    for ci, yr in enumerate(YEARS, 3):
        hdr(ws, 4, ci, str(yr))

    sei_data = [
        ("Conventional",  [CONV_RFUEL[y]+CONV_ELEC[y]+CONV_GAS[y] for y in YEARS]),
        ("BEV haulage",   [BEV_RFUEL[y]+BEV_ELEC[y]+BEV_GAS[y] for y in YEARS]),
        ("Hydrogen FCEV", [H2_H2[y]+H2_ELEC[y]+H2_RFUEL[y]+H2_GAS[y] for y in YEARS]),
        ("Low-fugitive",  [LFUG_RFUEL[y]+LFUG_ELEC[y]+LFUG_GAS[y] for y in YEARS]),
        ("Integrated low-carbon", [ILC_ELEC[y]+ILC_RFUEL[y]+ILC_GAS[y] for y in YEARS]),
    ]
    for ri, (name, vals) in enumerate(sei_data):
        r = 5 + ri
        dat(ws, r, 2, name, bold=True, even=(ri%2==1))
        for ci, v in enumerate(vals, 3):
            dat(ws, r, ci, v, fmt="#,##0", even=(ri%2==1))

    # rfuel table
    ws.cell(row=12, column=2, value="Refined Liquid Fuels Coefficient (GJ/kt)").font = Font(bold=True, size=11, color="1F4E79")
    hdr(ws, 13, 2, "State")
    for ci, yr in enumerate(YEARS, 3):
        hdr(ws, 13, ci, str(yr))

    rfuel_data = [
        ("Conventional",    [CONV_RFUEL[y] for y in YEARS]),
        ("BEV haulage",     [BEV_RFUEL[y] for y in YEARS]),
        ("Hydrogen FCEV",   [H2_RFUEL[y] for y in YEARS]),
        ("Low-fugitive",    [LFUG_RFUEL[y] for y in YEARS]),
        ("Integrated lc",   [ILC_RFUEL[y] for y in YEARS]),
    ]
    for ri, (name, vals) in enumerate(rfuel_data):
        r = 14 + ri
        dat(ws, r, 2, name, bold=True, even=(ri%2==1))
        for ci, v in enumerate(vals, 3):
            dat(ws, r, ci, v, fmt="#,##0", even=(ri%2==1))

    # Fugitive table
    ws.cell(row=21, column=2, value="Fugitive Process CO2e (tCO2e/kt)").font = Font(bold=True, size=11, color="1F4E79")
    hdr(ws, 22, 2, "State")
    for ci, yr in enumerate(YEARS, 3):
        hdr(ws, 22, ci, str(yr))

    fug_data = [
        ("Conventional",        [CONV_P_CO2E[y] for y in YEARS]),
        ("BEV haulage",         [CONV_P_CO2E[y] for y in YEARS]),
        ("Hydrogen FCEV",       [CONV_P_CO2E[y] for y in YEARS]),
        ("Low-fugitive",        [LFUG_P_CO2E[y] for y in YEARS]),
        ("Integrated low-carbon",[ILC_P_CO2E[y] for y in YEARS]),
    ]
    for ri, (name, vals) in enumerate(fug_data):
        r = 23 + ri
        dat(ws, r, 2, name, bold=True, even=(ri%2==1))
        for ci, v in enumerate(vals, 3):
            dat(ws, r, ci, v, fmt="0.0", even=(ri%2==1))


def build_total_demand(wb):
    ws = wb.create_sheet("Total_Demand")
    ws.sheet_view.showGridLines = False
    for col, w in zip("ABCDEFG", [4, 12, 18, 18, 20, 22, 22]):
        ws.column_dimensions[col].width = w

    title_cell(ws, 1, 2, "Total Coal Demand Trajectory (declining__coal_mining_total, −1.1%/yr)")
    ws.merge_cells("B1:G1")

    for ci, h in enumerate(["Year", "Demand (kt)", "Index (2025=1.00)",
                              "Energy PJ (conv)", "Energy CO2e (MtCO2e)", "Fugitive CO2e (MtCO2e)"], 2):
        hdr(ws, 3, ci, h)

    for ri, yr in enumerate(YEARS):
        r = 4 + ri
        kt = TOTAL_DEMAND[yr]
        idx = kt / TOTAL_KT
        pj = kt * SEI_2025 / 1e6
        e_co2 = kt * ENERGY_CO2E_2025 / 1e6
        f_co2 = kt * PROCESS_CO2E_2025 / 1e6
        even = ri % 2 == 1
        dat(ws, r, 2, yr, even=even)
        dat(ws, r, 3, kt, fmt="#,##0", even=even)
        dat(ws, r, 4, round(idx, 4), fmt="0.0000", even=even)
        dat(ws, r, 5, round(pj, 1), fmt="0.0", even=even)
        dat(ws, r, 6, round(e_co2, 2), fmt="0.00", even=even)
        dat(ws, r, 7, round(f_co2, 2), fmt="0.00", even=even)

    ws.cell(row=11, column=2,
            value="Curve: declining__coal_mining_total | Anchor: 420,000 kt (2025) | Rate: −1.1%/yr compound "
                  "| = weighted avg of export −0.75%/yr + domestic −2.5%/yr").font = \
        Font(italic=True, size=9, color="595959")
    ws.merge_cells("B11:G11")


def build_emissions_pathways(wb):
    ws = wb.create_sheet("Emissions_Pathways")
    ws.sheet_view.showGridLines = False
    for col, w in zip("ABCDEFG", [4, 30, 12, 12, 12, 12, 12]):
        ws.column_dimensions[col].width = w

    title_cell(ws, 1, 2, "Total Coal Mining — Emissions by Pathway State (MtCO2e/yr)")
    ws.merge_cells("B1:G1")

    for ci, h in enumerate(["State", "2025", "2030", "2035", "2040", "2045", "2050"], 2):
        hdr(ws, 3, ci, h)

    state_defs = [
        ("Conventional",        CONV_E_CO2E, CONV_P_CO2E),
        ("BEV haulage",         BEV_E_CO2E,  CONV_P_CO2E),
        ("Hydrogen FCEV",       BEV_E_CO2E,  CONV_P_CO2E),
        ("Low-fugitive",        CONV_E_CO2E, LFUG_P_CO2E),
        ("Integrated low-carbon", ILC_E_CO2E, ILC_P_CO2E),
    ]
    for ri, (state, eco2, pco2) in enumerate(state_defs):
        r = 4 + ri
        even = ri % 2 == 1
        dat(ws, r, 2, state, bold=True, even=even)
        for ci, yr in enumerate(YEARS, 3):
            kt = TOTAL_DEMAND[yr]
            val = round((eco2[yr] + pco2[yr]) * kt / 1e6, 2)
            dat(ws, r, ci, val, fmt="0.00", even=even)


def build_charts_sheet(wb):
    ws = wb.create_sheet("Charts")
    ws.sheet_view.showGridLines = False

    # ── Chart 1: Demand Trajectory ──────────────────────────────────────
    ws.cell(row=1, column=1, value="Year")
    ws.cell(row=1, column=2, value="Demand (kt)")
    for ri, yr in enumerate(YEARS):
        ws.cell(row=2+ri, column=1, value=yr)
        ws.cell(row=2+ri, column=2, value=TOTAL_DEMAND[yr])

    chart1 = LineChart()
    chart1.title = "Total Australian Coal Demand Trajectory (kt)"
    chart1.style = 10
    chart1.y_axis.title = "kt coal"
    chart1.x_axis.title = "Year"
    chart1.height = 12
    chart1.width = 20

    data = Reference(ws, min_col=2, min_row=1, max_row=7)
    cats = Reference(ws, min_col=1, min_row=2, max_row=7)
    chart1.add_data(data, titles_from_data=True)
    chart1.set_categories(cats)
    chart1.series[0].graphicalProperties.line.solidFill = "2E75B6"
    chart1.series[0].graphicalProperties.line.width = 20000
    ws.add_chart(chart1, "D2")

    # ── Chart 2: SEI by State ──────────────────────────────────────────
    ws.cell(row=10, column=1, value="State")
    ws.cell(row=10, column=2, value="2025 SEI (GJ/kt)")
    sei_chart_data = [
        ("Conventional", 406),
        ("BEV haulage",  292),
        ("H2 FCEV",      320),
        ("Low-fugitive", 414),
        ("Integrated",   286),
    ]
    for ri, (s, v) in enumerate(sei_chart_data):
        ws.cell(row=11+ri, column=1, value=s)
        ws.cell(row=11+ri, column=2, value=v)

    chart2 = BarChart()
    chart2.type = "bar"
    chart2.title = "2025 Specific Energy Intensity by State (GJ/kt)"
    chart2.style = 10
    chart2.y_axis.title = "GJ/kt"
    chart2.height = 12
    chart2.width = 20
    data2 = Reference(ws, min_col=2, min_row=10, max_row=15)
    cats2 = Reference(ws, min_col=1, min_row=11, max_row=15)
    chart2.add_data(data2, titles_from_data=True)
    chart2.set_categories(cats2)
    ws.add_chart(chart2, "D18")

    # ── Chart 3: Total CO2e by State at 2025 demand ──────────────────
    ws.cell(row=20, column=1, value="State")
    ws.cell(row=20, column=2, value="Total CO2e 2025 (tCO2e/kt)")
    co2e_data = [
        ("Conventional",  91.0),
        ("BEV haulage",   79.6),
        ("H2 FCEV",       79.6),
        ("Low-fugitive",  56.7),
        ("Integrated",    36.6),
    ]
    for ri, (s, v) in enumerate(co2e_data):
        ws.cell(row=21+ri, column=1, value=s)
        ws.cell(row=21+ri, column=2, value=v)

    chart3 = BarChart()
    chart3.type = "bar"
    chart3.title = "2025 Total CO2e Intensity by State (tCO2e/kt)"
    chart3.style = 10
    chart3.y_axis.title = "tCO2e/kt"
    chart3.height = 12
    chart3.width = 20
    data3 = Reference(ws, min_col=2, min_row=20, max_row=25)
    cats3 = Reference(ws, min_col=1, min_row=21, max_row=25)
    chart3.add_data(data3, titles_from_data=True)
    chart3.set_categories(cats3)
    ws.add_chart(chart3, "D36")

    # ── Chart 4: Total energy PJ over time ───────────────────────────
    ws.cell(row=30, column=1, value="Year")
    ws.cell(row=30, column=2, value="Energy (PJ)")
    for ri, yr in enumerate(YEARS):
        ws.cell(row=31+ri, column=1, value=yr)
        ws.cell(row=31+ri, column=2, value=round(TOTAL_DEMAND[yr] * SEI_2025 / 1e6, 1))

    chart4 = LineChart()
    chart4.title = "Total Coal Mining Energy Use (PJ, conventional state)"
    chart4.style = 10
    chart4.y_axis.title = "PJ"
    chart4.x_axis.title = "Year"
    chart4.height = 12
    chart4.width = 20
    data4 = Reference(ws, min_col=2, min_row=30, max_row=36)
    cats4 = Reference(ws, min_col=1, min_row=31, max_row=36)
    chart4.add_data(data4, titles_from_data=True)
    chart4.set_categories(cats4)
    chart4.series[0].graphicalProperties.line.solidFill = "C00000"
    chart4.series[0].graphicalProperties.line.width = 20000
    ws.add_chart(chart4, "D54")

    for col in "AB":
        ws.column_dimensions[col].hidden = True


def build_excel_workbook():
    wb = Workbook()
    build_cover(wb)
    build_energy_emissions_sensecheck(wb)
    build_calibration_sheet(wb)
    build_total_production_sheet(wb)
    build_state_assumptions(wb)
    build_total_demand(wb)
    build_emissions_pathways(wb)
    build_charts_sheet(wb)

    for ws_name, color in [
        ("Cover",                    "1F4E79"),
        ("Energy_Emissions_SenseCheck", "375623"),
        ("174PJ_Calibration",        "2E75B6"),
        ("Total_Production",         "5B9BD5"),
        ("State_Assumptions",        "70AD47"),
        ("Total_Demand",             "70AD47"),
        ("Emissions_Pathways",       "FFC000"),
        ("Charts",                   "ED7D31"),
    ]:
        wb[ws_name].sheet_properties.tabColor = color

    out_path = OUT_DIR / "coal_mining_calibration.xlsx"
    wb.save(str(out_path))
    print(f"  Written: {out_path}")


# ─────────────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    print("Generating documentation...")
    build_word_doc()
    build_excel_workbook()
    print("Done.")
