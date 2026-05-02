#!/usr/bin/env python3
"""
Create calibration Excel workbooks and Word reports for all 6 new sector families.
Run from: c:/code/simple-msm
"""

import os
import datetime

# ── openpyxl ─────────────────────────────────────────────────────────────────
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

# ── python-docx ──────────────────────────────────────────────────────────────
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE = os.path.dirname(os.path.abspath(__file__))
TODAY = datetime.date.today().isoformat()

# ── colour palette ───────────────────────────────────────────────────────────
NAVY_FILL  = PatternFill("solid", fgColor="1F3864")
BLUE_FILL  = PatternFill("solid", fgColor="B8CCE4")
WHITE_FONT = Font(color="FFFFFF", bold=True, size=11)
BOLD_FONT  = Font(bold=True, size=10)
BODY_FONT  = Font(size=10)
HEADER_FONT = Font(bold=True, size=10)

thin_border = Border(
    left=Side(style="thin"), right=Side(style="thin"),
    top=Side(style="thin"), bottom=Side(style="thin")
)


def style_title_row(ws, row, ncols):
    for col in range(1, ncols + 1):
        c = ws.cell(row=row, column=col)
        c.fill = NAVY_FILL
        c.font = WHITE_FONT
        c.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)


def style_header_row(ws, row, ncols):
    for col in range(1, ncols + 1):
        c = ws.cell(row=row, column=col)
        c.fill = BLUE_FILL
        c.font = HEADER_FONT
        c.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
        c.border = thin_border


def style_data_row(ws, row, ncols):
    for col in range(1, ncols + 1):
        c = ws.cell(row=row, column=col)
        c.font = BODY_FONT
        c.alignment = Alignment(vertical="top", wrap_text=True)
        c.border = thin_border


# ═══════════════════════════════════════════════════════════════════════════════
#  FAMILY METADATA
# ═══════════════════════════════════════════════════════════════════════════════

FAMILIES = [
    {
        "id":     "domestic_aviation",
        "label":  "Domestic Aviation",
        "unit":   "million_pkm",
        "anchor": 75_500,
        "aes_pj": 85.0,
        "nggi_mt": 6.1,
        "nggi_cat": "Cat 1A3a domestic aviation",
        "intensity_2025": 1126.0,
        "fuels": ["aviation_turbine_fuel"],
        "coeff_2025": [1126.0],
        "ef": [71.5],
        "sources": ["S001 — AES 2023-24 Table F", "S004 — National Inventory Report 2023", "S012 — BITRE Aviation Statistical Report 2023-24"],
        "states": [
            ("domestic_aviation__conventional_jet", "Conventional jet", "incumbent", 1.00, 0.60, 0.05, "1126.0 → 1000.0 GJ/million_pkm"),
            ("domestic_aviation__saf_blend", "SAF blend", "progression", 0.02, 0.60, 0.08, "ATF+SAF blend 5%→100% SAF"),
            ("domestic_aviation__electric_short_haul", "Battery-electric short-haul", "progression", 0.00, 0.20, 0.12, "360.0 → 280.0 GJ/million_pkm"),
        ],
        "key_assumptions": [
            "Demand anchor: 75,500 million_pkm (BITRE Aviation Statistical Report 2023-24)",
            "Energy intensity: 85 PJ / 75,500 million_pkm = 1,126 GJ/million_pkm (AES 2023-24)",
            "100% aviation turbine fuel in incumbent state",
            "SAF biogenic CO2 excluded from scope 1 per NGA accounting",
            "Electricity-related emissions excluded per A003 (scope 1 only)",
            "GWP: IPCC AR5 GWP100",
        ],
        "caveats": [
            "National average masks route-length variation (short-haul vs long-haul domestic)",
            "SAF supply chain for Australia is nascent; 100% SAF by 2050 is ambitious",
            "Electric aviation limited to <500 km routes; max share 20% reflects route constraints",
            "NGGI Cat 1A3a total 6.1 MtCO2e assumes all domestic fuel combustion in scope",
        ],
    },
    {
        "id":     "international_aviation",
        "label":  "International Aviation",
        "unit":   "million_pkm",
        "anchor": 70_000,
        "aes_pj": 165.0,
        "nggi_mt": 11.8,
        "nggi_cat": "International aviation bunkers",
        "intensity_2025": 2357.0,
        "fuels": ["aviation_turbine_fuel"],
        "coeff_2025": [2357.0],
        "ef": [71.5],
        "sources": ["S001 — AES 2023-24 international aviation bunkers", "S004 — National Inventory Report 2023", "S012 — BITRE international air travel 2023-24"],
        "states": [
            ("international_aviation__conventional_jet", "Conventional jet", "incumbent", 1.00, 0.70, 0.04, "2357.0 → 2100.0 GJ/million_pkm"),
            ("international_aviation__saf_blend", "SAF blend", "progression", 0.00, 0.50, 0.07, "ATF 100%→20% + SAF 0%→80%"),
            ("international_aviation__hydrogen_aircraft", "Hydrogen aircraft", "progression", 0.00, 0.15, 0.15, "0 → 2200 → 2000 GJ H2/million_pkm (from 2035)"),
        ],
        "key_assumptions": [
            "Demand anchor: 70,000 million_pkm (BITRE international air travel 2023-24)",
            "Energy intensity: 165 PJ / 70,000 million_pkm = 2,357 GJ/million_pkm",
            "International aviation bunkers (fuel uplifted in Australia)",
            "SAF biogenic CO2 excluded from scope 1 per NGA",
            "Hydrogen aircraft available from 2035 only",
            "GWP: IPCC AR5 GWP100",
        ],
        "caveats": [
            "International bunker allocation has boundary uncertainty (vs ICAO fuel uplift data)",
            "SAF penetration slower than domestic due to supply constraints for long-haul",
            "H2 aircraft technology is highly uncertain; max 15% by 2050 is exploratory",
            "Methane slip for LNG not modelled in Phase 1",
        ],
    },
    {
        "id":     "domestic_shipping",
        "label":  "Domestic Shipping",
        "unit":   "million_tkm",
        "anchor": 30_000,
        "aes_pj": 30.0,
        "nggi_mt": 2.28,
        "nggi_cat": "Domestic navigation",
        "intensity_2025": 1000.0,
        "fuels": ["marine_diesel_oil", "heavy_fuel_oil"],
        "coeff_2025": [650.0, 350.0],
        "ef": [74.4, 78.9],
        "sources": ["S001 — AES 2023-24 Table F (ANZSIC 4800)", "S004 — National Inventory Report 2023", "S013 — BITRE coastal freight statistics"],
        "states": [
            ("domestic_shipping__conventional_diesel", "Conventional diesel", "incumbent", 1.00, 0.65, 2.5, "MDO 650 → 580 + HFO 350 → 310 GJ/million_tkm"),
            ("domestic_shipping__battery_electric_vessel", "Battery-electric vessel", "progression", 0.01, 0.25, 5.0, "450 → 380 GJ electricity/million_tkm"),
            ("domestic_shipping__green_ammonia_vessel", "Green ammonia vessel", "progression", 0.00, 0.20, 8.0, "1200 → 1000 GJ ammonia/million_tkm"),
        ],
        "key_assumptions": [
            "Demand anchor: 30,000 million_tkm = 30 billion tkm (BITRE coastal freight)",
            "Energy intensity: 30 PJ / 30,000 million_tkm = 1,000 GJ/million_tkm",
            "Fuel split: 65% MDO, 35% HFO (vessel type mix)",
            "Green ammonia combustion CO2 = 0 (N2O slip excluded Phase 1)",
            "Electric vessel scope 1 CO2 = 0; electricity scope 2 excluded per A003",
            "GWP: IPCC AR5 GWP100",
        ],
        "caveats": [
            "Coastal navigation is heterogeneous; 65/35 MDO/HFO split is an approximation",
            "N2O emissions from ammonia combustion excluded in Phase 1 (add in Phase 2)",
            "Battery-electric range limits applicability to <100 km coastal routes",
            "Green ammonia supply chain is nascent; max 20% by 2050 is exploratory",
        ],
    },
    {
        "id":     "international_shipping",
        "label":  "International Shipping",
        "unit":   "million_tkm",
        "anchor": 200_000,
        "aes_pj": 65.0,
        "nggi_mt": 5.0,
        "nggi_cat": "International maritime bunkers",
        "intensity_2025": 325.0,
        "fuels": ["heavy_fuel_oil", "marine_diesel_oil"],
        "coeff_2025": [195.0, 130.0],
        "ef": [78.9, 74.4],
        "sources": ["S001 — AES 2023-24 international maritime bunkers", "S004 — National Inventory Report 2023", "S013 — BITRE freight statistics"],
        "states": [
            ("international_shipping__conventional_hfo", "Conventional HFO/MDO", "incumbent", 1.00, 0.55, 1.5, "HFO 195 → 165 + MDO 130 → 110 GJ/million_tkm"),
            ("international_shipping__lng_transition", "LNG transition", "progression", 0.01, 0.30, 1.8, "350 → 290 GJ LNG/million_tkm"),
            ("international_shipping__zero_emission_fuel", "Zero-emission fuel (ammonia)", "progression", 0.00, 0.25, 4.0, "0 → 400 → 320 GJ ammonia/million_tkm (from 2030)"),
        ],
        "key_assumptions": [
            "Service volume 200,000 million_tkm is estimated (no direct Australian statistic)",
            "Energy intensity: 65 PJ / 200,000 million_tkm = 325 GJ/million_tkm",
            "Fuel split: 60% HFO, 40% MDO for international large vessels",
            "Zero-emission fuel state available from 2030 only",
            "Methane slip from LNG excluded in Phase 1",
            "GWP: IPCC AR5 GWP100",
        ],
        "caveats": [
            "Service volume (200,000 million_tkm) is estimated; cross-check with IMO GHG data in Phase 2",
            "LNG methane slip can increase net GHG relative to HFO; excluded Phase 1",
            "N2O from ammonia combustion excluded Phase 1",
            "International bunker data has flag/uplift-location boundary uncertainty",
        ],
    },
    {
        "id":     "rail_passenger",
        "label":  "Rail Passenger",
        "unit":   "million_pkm",
        "anchor": 22_000,
        "aes_pj": 8.0,
        "nggi_mt": 0.27,
        "nggi_cat": "Rail passenger scope 1 diesel",
        "intensity_2025": 364.0,
        "fuels": ["electricity", "diesel"],
        "coeff_2025": [182.0, 182.0],
        "ef": [0.0, 69.9],
        "sources": ["S001 — AES 2023-24 estimated rail passenger share", "S004 — National Inventory Report 2023", "S012 — BITRE Rail Summary Data"],
        "states": [
            ("rail_passenger__conventional_mixed", "Conventional mixed (electric+diesel)", "incumbent", 1.00, 0.55, 0.08, "Elec 182 → 195 + Diesel 182 → 120 GJ/million_pkm"),
            ("rail_passenger__fully_electrified", "Fully electrified OHL", "progression", 0.15, 0.60, 0.07, "250 → 200 GJ electricity/million_pkm"),
            ("rail_passenger__hydrogen_regional", "Hydrogen regional train", "progression", 0.00, 0.20, 0.15, "0 → 320 → 250 GJ H2/million_pkm (from 2030)"),
            ("rail_passenger__battery_electric_regional", "Battery-electric regional", "progression", 0.01, 0.15, 0.12, "280 → 230 GJ electricity/million_pkm"),
        ],
        "key_assumptions": [
            "Demand anchor: 22,000 million_pkm (BITRE Rail Summary Data)",
            "AES rail passenger energy estimated at ~8 PJ (total rail ~33 PJ; passenger share ~24%)",
            "National average 50% electric / 50% diesel by energy (urban electric, regional diesel)",
            "Scope 1 CO2e from diesel only; electricity scope 2 excluded per A003",
            "Hydrogen trains available from 2030; battery-electric available from 2025",
            "GWP: IPCC AR5 GWP100",
        ],
        "caveats": [
            "AES does not separately report freight vs passenger rail; 8 PJ is an estimate",
            "50/50 energy split is a national average; state-level variation is significant",
            "OHL electrification requires 20-30 year investment lead times",
            "Battery-electric limited to <200 km per charge; hydrogen limited by H2 supply chain",
        ],
    },
    {
        "id":     "rail_freight",
        "label":  "Rail Freight",
        "unit":   "billion_tkm",
        "anchor": 700,
        "aes_pj": 25.0,
        "nggi_mt": 1.62,
        "nggi_cat": "Rail freight scope 1 diesel",
        "intensity_2025": 35715.0,
        "fuels": ["diesel", "electricity"],
        "coeff_2025": [33929.0, 1786.0],
        "ef": [69.9, 0.0],
        "sources": ["S001 — AES 2023-24 estimated rail freight share", "S004 — National Inventory Report 2023", "S013 — BITRE freight linehaul statistics"],
        "states": [
            ("rail_freight__diesel_electric", "Diesel-electric locomotive", "incumbent", 1.00, 0.60, 0.015, "Diesel 33929 → 28500 + Elec 1786 → 1500 GJ/billion_tkm"),
            ("rail_freight__overhead_electrification", "Overhead electrification", "progression", 0.01, 0.30, 0.020, "25000 → 20000 GJ electricity/billion_tkm"),
            ("rail_freight__hydrogen_locomotive", "Hydrogen locomotive", "progression", 0.00, 0.20, 0.04, "0 → 40000 → 30000 GJ H2/billion_tkm (from 2030)"),
        ],
        "key_assumptions": [
            "Demand anchor: 700 billion_tkm (BITRE freight linehaul; mainly iron ore and coal)",
            "AES rail freight energy estimated at ~25 PJ (total rail ~33 PJ; freight share ~76%)",
            "95% diesel, 5% electricity (electrified corridor fraction)",
            "Scope 1 CO2e from diesel only; electricity scope 2 excluded per A003",
            "Hydrogen locomotives available from 2030 only",
            "GWP: IPCC AR5 GWP100",
        ],
        "caveats": [
            "AES does not separately report freight vs passenger rail; 25 PJ is an estimate",
            "Heavy-haul iron ore (Pilbara private rail) dominates volumes but is separately operated",
            "OHL electrification is very capital-intensive (AUD 1-3M/km); limited to high-density corridors",
            "H2 locomotive technology is at demonstration scale; max 20% by 2050 is exploratory",
        ],
    },
]


# ═══════════════════════════════════════════════════════════════════════════════
#  EXCEL WORKBOOK CREATION
# ═══════════════════════════════════════════════════════════════════════════════

def create_excel_workbook(fam):
    wb = Workbook()
    wb.remove(wb.active)

    # ── 1. Summary ─────────────────────────────────────────────────────────
    ws = wb.create_sheet("Summary")
    ws.column_dimensions["A"].width = 30
    ws.column_dimensions["B"].width = 55

    # Title
    ws.merge_cells("A1:B1")
    ws["A1"] = f"{fam['label']} — Phase 1 Calibration Workbook"
    ws["A1"].font = Font(bold=True, size=14, color="FFFFFF")
    ws["A1"].fill = NAVY_FILL
    ws["A1"].alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 30

    rows = [
        ("Family ID", fam["id"]),
        ("Output unit", fam["unit"]),
        ("Calibration date", TODAY),
        ("Demand anchor 2025", f"{fam['anchor']:,} {fam['unit']}"),
        ("AES 2023-24 total energy", f"{fam['aes_pj']} PJ"),
        ("Energy intensity 2025", f"{fam['intensity_2025']:,.1f} GJ/{fam['unit']}"),
        ("NGGI total CO2e 2025", f"{fam['nggi_mt']} MtCO2e ({fam['nggi_cat']})"),
        ("Number of states", str(len(fam["states"]))),
        ("Scope boundary", "Scope 1 direct combustion only"),
        ("GWP basis", "IPCC AR5 GWP100"),
        ("Currency", "AUD_2024 (real 2024 AUD)"),
    ]

    for i, (k, v) in enumerate(rows, start=2):
        ws.cell(row=i, column=1, value=k).font = BOLD_FONT
        ws.cell(row=i, column=2, value=v).font = BODY_FONT
        if i % 2 == 0:
            ws.cell(row=i, column=1).fill = PatternFill("solid", fgColor="EBF3FB")
            ws.cell(row=i, column=2).fill = PatternFill("solid", fgColor="EBF3FB")

    # ── 2. AES Calibration ─────────────────────────────────────────────────
    ws2 = wb.create_sheet("AES_Calibration")
    ws2.column_dimensions["A"].width = 28
    ws2.column_dimensions["B"].width = 18
    ws2.column_dimensions["C"].width = 20
    ws2.column_dimensions["D"].width = 18
    ws2.column_dimensions["E"].width = 20
    ws2.column_dimensions["F"].width = 14

    ws2.merge_cells("A1:F1")
    ws2["A1"] = f"{fam['label']} — AES Calibration Check"
    ws2["A1"].font = Font(bold=True, size=12, color="FFFFFF")
    ws2["A1"].fill = NAVY_FILL
    ws2["A1"].alignment = Alignment(horizontal="center")

    hdrs = ["Fuel type", "AES 2023-24 share", "Coefficient (GJ/unit)", "Anchor quantity", "Model total (PJ)", "Coverage %"]
    for col, h in enumerate(hdrs, 1):
        c = ws2.cell(row=2, column=col, value=h)
        c.fill = BLUE_FILL
        c.font = HEADER_FONT
        c.border = thin_border

    # Compute coverage
    anchor = fam["anchor"]
    total_model_pj = sum(c * anchor / 1e6 for c in fam["coeff_2025"])
    coverage = total_model_pj / fam["aes_pj"] * 100

    for i, (fuel, coeff, ef) in enumerate(zip(fam["fuels"], fam["coeff_2025"], fam["ef"]), 3):
        fuel_pj = coeff * anchor / 1e6
        fuel_share_pj = fam["aes_pj"] * (coeff / sum(fam["coeff_2025"])) if sum(fam["coeff_2025"]) > 0 else 0
        row_vals = [fuel, f"{fuel_share_pj:.1f} PJ", f"{coeff:,.1f}", f"{anchor:,}", f"{fuel_pj:.2f}", f"{fuel_pj/fam['aes_pj']*100:.1f}%"]
        for col, v in enumerate(row_vals, 1):
            c = ws2.cell(row=i, column=col, value=v)
            c.font = BODY_FONT
            c.border = thin_border

    # Summary row
    r = len(fam["fuels"]) + 3
    ws2.cell(row=r, column=1, value="TOTAL").font = BOLD_FONT
    ws2.cell(row=r, column=2, value=f"{fam['aes_pj']:.1f} PJ (AES reference)").font = BOLD_FONT
    ws2.cell(row=r, column=5, value=f"{total_model_pj:.2f}").font = BOLD_FONT
    ws2.cell(row=r, column=6, value=f"{coverage:.1f}%").font = BOLD_FONT

    # ── 3. State Parameters ────────────────────────────────────────────────
    ws3 = wb.create_sheet("State_Parameters")
    for col, w in enumerate([32, 12, 12, 35, 14, 14], 1):
        ws3.column_dimensions[get_column_letter(col)].width = w

    ws3.merge_cells("A1:F1")
    ws3["A1"] = f"{fam['label']} — State Parameters by Year"
    ws3["A1"].font = Font(bold=True, size=12, color="FFFFFF")
    ws3["A1"].fill = NAVY_FILL
    ws3["A1"].alignment = Alignment(horizontal="center")

    hdrs3 = ["State ID", "Stage", "Max share 2025", "Coefficient 2025", "Max share 2050", "Cost 2025 (AUD/unit)"]
    for col, h in enumerate(hdrs3, 1):
        c = ws3.cell(row=2, column=col, value=h)
        c.fill = BLUE_FILL
        c.font = HEADER_FONT
        c.border = thin_border

    for i, (sid, slabel, stage, ms25, ms50, cost25, coeff_desc) in enumerate(fam["states"], 3):
        vals = [sid, stage, ms25, coeff_desc, ms50, cost25]
        for col, v in enumerate(vals, 1):
            c = ws3.cell(row=i, column=col, value=v)
            c.font = BODY_FONT
            c.border = thin_border

    # ── 4. Emissions ───────────────────────────────────────────────────────
    ws4 = wb.create_sheet("Emissions")
    for col, w in enumerate([32, 14, 14, 14, 20], 1):
        ws4.column_dimensions[get_column_letter(col)].width = w

    ws4.merge_cells("A1:E1")
    ws4["A1"] = f"{fam['label']} — Emissions by State"
    ws4["A1"].font = Font(bold=True, size=12, color="FFFFFF")
    ws4["A1"].fill = NAVY_FILL
    ws4["A1"].alignment = Alignment(horizontal="center")

    hdrs4 = ["State ID", "Energy CO2e 2025", "Process CO2e 2025", "Total CO2e (Mt/yr)", "Boundary"]
    for col, h in enumerate(hdrs4, 1):
        c = ws4.cell(row=2, column=col, value=h)
        c.fill = BLUE_FILL
        c.font = HEADER_FONT
        c.border = thin_border

    for i, (sid, slabel, stage, ms25, ms50, cost25, coeff_desc) in enumerate(fam["states"], 3):
        # Compute energy CO2e for first state only (incumbent)
        if i == 3 and fam["coeff_2025"] and fam["ef"]:
            eco2e = sum(c * e for c, e in zip(fam["coeff_2025"], fam["ef"])) / 1000
            total_mt = round(eco2e * fam["anchor"] / 1e6, 3)
        else:
            eco2e = "—"
            total_mt = "—"
        vals = [sid, eco2e if i == 3 else "0 (zero-carbon fuel or not-2025)", 0, total_mt if i == 3 else "—", "Scope 1 only"]
        for col, v in enumerate(vals, 1):
            c = ws4.cell(row=i, column=col, value=v)
            c.font = BODY_FONT
            c.border = thin_border

    # ── 5. Data Sources ────────────────────────────────────────────────────
    ws5 = wb.create_sheet("Data_Sources")
    ws5.column_dimensions["A"].width = 12
    ws5.column_dimensions["B"].width = 70

    ws5.merge_cells("A1:B1")
    ws5["A1"] = f"{fam['label']} — Data Sources"
    ws5["A1"].font = Font(bold=True, size=12, color="FFFFFF")
    ws5["A1"].fill = NAVY_FILL
    ws5["A1"].alignment = Alignment(horizontal="center")

    hdrs5 = ["Source ID", "Description"]
    for col, h in enumerate(hdrs5, 1):
        c = ws5.cell(row=2, column=col, value=h)
        c.fill = BLUE_FILL
        c.font = HEADER_FONT
        c.border = thin_border

    for i, src in enumerate(fam["sources"], 3):
        parts = src.split(" — ", 1)
        sid = parts[0].strip()
        sdesc = parts[1].strip() if len(parts) > 1 else ""
        ws5.cell(row=i, column=1, value=sid).font = BODY_FONT
        ws5.cell(row=i, column=2, value=sdesc).font = BODY_FONT
        ws5.cell(row=i, column=1).border = thin_border
        ws5.cell(row=i, column=2).border = thin_border

    # Also list assumptions
    ws5.cell(row=len(fam["sources"]) + 4, column=1, value="Assumptions").font = BOLD_FONT
    for i, asm in enumerate(["A002 — cost definition excludes commodities", "A003 — scope 1 boundary", "A022 — feasibility bounds", "A023 — cost path smoothing"], len(fam["sources"]) + 5):
        parts = asm.split(" — ", 1)
        ws5.cell(row=i, column=1, value=parts[0]).font = BODY_FONT
        ws5.cell(row=i, column=2, value=parts[1] if len(parts) > 1 else "").font = BODY_FONT

    # Save
    dir_path = os.path.join(BASE, fam["id"])
    out_path = os.path.join(dir_path, f"{fam['id']}_calibration.xlsx")
    wb.save(out_path)
    print(f"  Written: {out_path}")


# ═══════════════════════════════════════════════════════════════════════════════
#  WORD REPORT CREATION
# ═══════════════════════════════════════════════════════════════════════════════

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.style.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p


def create_word_report(fam):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # Title
    title = doc.add_heading(f"{fam['label']} — Phase 1 Calibration Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(f"Family: {fam['id']}   |   Generated: {TODAY}   |   Unit: {fam['unit']}").italic = True

    doc.add_paragraph()

    # 1. Calibration Basis
    add_heading(doc, "1. Calibration Basis", 1)
    p = doc.add_paragraph()
    p.add_run("Demand anchor. ").bold = True
    p.add_run(f"The 2025 baseline demand is anchored at {fam['anchor']:,} {fam['unit']}.")

    p2 = doc.add_paragraph()
    p2.add_run("Energy calibration. ").bold = True
    p2.add_run(
        f"Total final energy at the 2025 anchor is {fam['aes_pj']} PJ (AES 2023-24). "
        f"This gives an energy intensity of {fam['intensity_2025']:,.1f} GJ/{fam['unit']}. "
        f"Input fuels: {', '.join(fam['fuels'])}."
    )

    p3 = doc.add_paragraph()
    p3.add_run("Emissions calibration. ").bold = True
    p3.add_run(
        f"Scope 1 CO2e at the anchor equals {fam['nggi_mt']} MtCO2e ({fam['nggi_cat']}). "
        f"This uses NGA 2025 emission factors and IPCC AR5 GWP100. "
        f"Electricity-related emissions are excluded (scope 2 boundary, per A003)."
    )

    # 2. Technology States
    add_heading(doc, "2. Technology States", 1)
    for (sid, slabel, stage, ms25, ms50, cost25, coeff_desc) in fam["states"]:
        add_heading(doc, slabel, 2)
        doc.add_paragraph(
            f"State ID: {sid}. Classification: {stage}. "
            f"Max share range: {ms25:.0%} (2025) to {ms50:.0%} (2050). "
            f"Cost: AUD {cost25}/{fam['unit']} (2024 real). "
            f"Energy trajectory: {coeff_desc}."
        )

    # 3. Key Assumptions
    add_heading(doc, "3. Key Assumptions", 1)
    for asm in fam["key_assumptions"]:
        add_bullet(doc, asm)

    # 4. Known Caveats
    add_heading(doc, "4. Known Caveats", 1)
    for cav in fam["caveats"]:
        add_bullet(doc, cav)

    # 5. Data Sources
    add_heading(doc, "5. Data Sources", 1)
    for src in fam["sources"]:
        add_bullet(doc, src)
    add_bullet(doc, "A002 — Cost definition: non-commodity conversion cost only")
    add_bullet(doc, "A003 — Scope 1 boundary: electricity-related emissions excluded")
    add_bullet(doc, "A022 — Feasibility bounds: indicative upper envelopes, not diffusion targets")
    add_bullet(doc, "A023 — Cost trajectories smoothed where evidence is sparse")

    # Save
    dir_path = os.path.join(BASE, fam["id"])
    out_path = os.path.join(dir_path, f"{fam['id']}_calibration_report.docx")
    doc.save(out_path)
    print(f"  Written: {out_path}")


# ═══════════════════════════════════════════════════════════════════════════════
#  MAIN
# ═══════════════════════════════════════════════════════════════════════════════
if __name__ == "__main__":
    print("\nCreating Excel calibration workbooks and Word reports...\n")
    for fam in FAMILIES:
        print(f"  Processing: {fam['label']} ({fam['id']})")
        create_excel_workbook(fam)
        create_word_report(fam)
        print()
    print("Done.")
