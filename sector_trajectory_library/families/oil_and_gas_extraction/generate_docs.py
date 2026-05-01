"""
Generate oil and gas extraction calibration documentation:
  - oil_and_gas_extraction_calibration_report.docx  (Word audit document)
  - oil_and_gas_extraction_calibration.xlsx         (Excel workbook)

Calibration basis:
  Total Australian gas production (6,100 PJ_gas, 2023-24, Geoscience Australia AECR 2025).
  Energy sense-check:  77,200 GJ/PJ × 6,100 PJ = 471.0 PJ  (AES 2023-24: 471 PJ, 100.1% ✓)
  Energy CO2e:          3,610 tCO2e/PJ × 6,100 PJ = 22.0 MtCO2e  (NGGI 2023-24 Cat 1A2 ✓)
  Fugitive CO2e:        3,610 tCO2e/PJ × 6,100 PJ = 22.0 MtCO2e  (NGGI 2023-24 Cat 1B2b ✓)

Discrepancy analysis vs previous model (gas_mining):
  1. Energy:        595 PJ → 471 PJ  (LNG liquefaction own-use excluded from AES final energy)
  2. Energy CO2e:  29.5 Mt → 22.0 Mt (wrong gas coefficient + flaring classified as fugitive)
  3. Fugitive CO2e: 20.0 Mt → 22.0 Mt (NGGI 2022-23 updated to NGGI 2023-24 with better measurement)
"""

from pathlib import Path
import docx
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import openpyxl
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.chart import BarChart, Reference
from openpyxl.utils import get_column_letter
import datetime

OUT_DIR = Path(__file__).parent

# ── Calibration constants ────────────────────────────────────────────────────
YEARS = [2025, 2030, 2035, 2040, 2045, 2050]
ANCHOR_PJ         = 6_100
AES_FINAL_ENERGY  = 471.0    # AES 2023-24 Table F (oil and gas extraction, ANZSIC 0700)
NGGI_ENERGY_MT    = 22.0     # NGGI 2023-24 Cat 1A2 (stationary energy — mining)
NGGI_FUGITIVE_MT  = 22.0     # NGGI 2023-24 Cat 1B2b (oil and gas fugitive)
MODELLED_PJ       = round(77_200 * ANCHOR_PJ / 1e6, 1)  # 471.0 PJ

# Previous (incorrect) model values for comparison
PREV_ENERGY_PJ    = 595.0
PREV_GAS_COEF     = 82_000
PREV_SEI          = 97_600
PREV_ENERGY_CO2E  = 4_840    # tCO2e/PJ
PREV_PROCESS_CO2E = 3_280    # tCO2e/PJ (NGGI 2022-23)

# Revised 2025 coefficients
GAS_2025   = 62_300  # GJ/PJ_gas
RFUEL_2025  =  9_800
ELEC_2025   =  5_100
SEI_2025    = 77_200

ENERGY_CO2E_2025  = 3_610   # tCO2e/PJ_gas (NGGI 2023-24 Cat 1A2 calibrated)
PROCESS_CO2E_2025 = 3_610   # tCO2e/PJ_gas (NGGI 2023-24 Cat 1B2b)

# AES-derived combustion CO2e (before NGGI calibration)
AES_DERIVED_CO2E_MT = (GAS_2025 * 51.4 + RFUEL_2025 * 69.9) / 1_000 * ANCHOR_PJ / 1e6  # 23.7 MtCO2e

# Demand trajectory: declining__oil_and_gas_extraction_total (-1.0%/yr)
TOTAL_DEMAND = {
    2025: 6_100,
    2030: 5_799,
    2035: 5_517,
    2040: 5_247,
    2045: 4_990,
    2050: 4_745,
}

# Conventional state trajectories (from family_states generator — source of truth)
CONV_GAS   = {2025: 62_300, 2030: 60_700, 2035: 59_200, 2040: 57_700, 2045: 56_500, 2050: 55_300}
CONV_RFUEL = {2025:  9_800, 2030:  9_600, 2035:  9_400, 2040:  9_200, 2045:  9_000, 2050:  8_800}
CONV_ELEC  = {2025:  5_100, 2030:  5_200, 2035:  5_300, 2040:  5_400, 2045:  5_500, 2050:  5_500}
CONV_E_CO2E= {2025: 3_610, 2030: 3_520, 2035: 3_440, 2040: 3_350, 2045: 3_280, 2050: 3_210}
CONV_P_CO2E= {2025: 3_610, 2030: 3_300, 2035: 3_000, 2040: 2_700, 2045: 2_450, 2050: 2_200}


# ── Style constants ──────────────────────────────────────────────────────────
DARK_BLUE  = "1F3864"
MID_BLUE   = "2E75B6"
LIGHT_BLUE = "D6E4F0"
GREEN_HDR  = "375623"
GREEN_LIGHT= "E2EFDA"
AMBER      = "FFC000"
RED        = "C00000"
ORANGE     = "ED7D31"
TEAL       = "00B0F0"
PURPLE     = "7030A0"


def thin_border():
    thin = Side(style="thin")
    return Border(left=thin, right=thin, top=thin, bottom=thin)

def hdr_fill(hex_color):
    return PatternFill("solid", fgColor=hex_color)

def cell_font(bold=False, color="000000", size=10):
    return Font(bold=bold, color=color, size=size)

def pct(val, ref):
    return f"{val/ref*100:.1f}%"

def set_col_widths(ws, widths_dict):
    for col_letter, width in widths_dict.items():
        ws.column_dimensions[col_letter].width = width


# ═══════════════════════════════════════════════════════════════════════════════
#  EXCEL WORKBOOK
# ═══════════════════════════════════════════════════════════════════════════════
def build_excel():
    wb = Workbook()
    wb.remove(wb.active)

    # ── Sheet 1: Energy_Emissions_SenseCheck ─────────────────────────────────
    ws1 = wb.create_sheet("Energy_Emissions_SenseCheck")
    ws1.sheet_properties.tabColor = GREEN_HDR

    ws1["A1"] = "Oil & Gas Extraction — Energy & Emissions Sense-Check (Conventional State, 2025–2050)"
    ws1["A1"].font = Font(bold=True, size=13, color=DARK_BLUE)
    ws1.merge_cells("A1:I1")
    ws1["A2"] = f"Generated: {datetime.date.today().isoformat()}  |  Calibration: AES 2023-24 Table F + NGGI 2023-24 Cat 1A2 & 1B2b"
    ws1["A2"].font = Font(italic=True, size=9, color="595959")

    # ── Discrepancy summary panel ─────────────────────────────────────────────
    ws1["A4"] = "CALIBRATION REVISION SUMMARY — Previous model (gas_mining) vs Revised model (oil_and_gas_extraction)"
    ws1["A4"].font = Font(bold=True, size=11, color=PURPLE)

    disc_headers = ["Issue", "Previous value", "Revised value", "Source of truth", "Root cause"]
    disc_data = [
        ["Final energy (PJ)",
         f"{PREV_ENERGY_PJ:.0f} PJ",
         f"{AES_FINAL_ENERGY:.0f} PJ",
         "AES 2023-24 Table F",
         "LNG liquefaction own-use (~316 PJ) incorrectly included. AES classifies liquefaction as 'energy transformation', not final energy."],
        ["Energy CO2e (MtCO2e)",
         f"{PREV_ENERGY_CO2E * ANCHOR_PJ / 1e6:.1f} MtCO2e",
         f"{ENERGY_CO2E_2025 * ANCHOR_PJ / 1e6:.1f} MtCO2e",
         "NGGI 2023-24 Cat 1A2",
         f"(a) Wrong gas coeff (82,000 → 62,300 GJ/PJ). (b) AES-derived combustion ({AES_DERIVED_CO2E_MT:.1f} MtCO2e) is 1.7 MtCO2e above NGGI 1A2 because routine flaring/venting is classified as fugitive (1B2b), not energy (1A2)."],
        ["Fugitive CO2e (MtCO2e)",
         f"{PREV_PROCESS_CO2E * ANCHOR_PJ / 1e6:.1f} MtCO2e",
         f"{PROCESS_CO2E_2025 * ANCHOR_PJ / 1e6:.1f} MtCO2e",
         "NGGI 2023-24 Cat 1B2b",
         "Previous model used NGGI 2022-23 (~20 MtCO2e). NGGI 2023-24 shows 22 MtCO2e, reflecting improved NGER direct-measurement methodology and real LNG fugitive growth."],
    ]

    for j, h in enumerate(disc_headers, 1):
        c = ws1.cell(row=5, column=j, value=h)
        c.fill = hdr_fill(PURPLE)
        c.font = cell_font(bold=True, color="FFFFFF", size=10)
        c.border = thin_border()
        c.alignment = Alignment(horizontal="center", wrap_text=True)
    ws1.row_dimensions[5].height = 28

    for i, row_data in enumerate(disc_data):
        for j, v in enumerate(row_data, 1):
            c = ws1.cell(row=6 + i, column=j, value=v)
            c.border = thin_border()
            c.alignment = Alignment(wrap_text=True, vertical="top")
            c.font = cell_font(size=9)
        ws1.row_dimensions[6 + i].height = 52

    set_col_widths(ws1, {"A": 24, "B": 20, "C": 20, "D": 24, "E": 60, "F": 16, "G": 16, "H": 16, "I": 14})
    disc_end_row = 9

    # ── Section 1: Fuel Consumption ───────────────────────────────────────────
    s1_start = disc_end_row + 2
    ws1.cell(row=s1_start, column=1, value="SECTION 1 — FUEL CONSUMPTION (conventional state, GJ/PJ_gas coefficient × total production PJ)").font = Font(bold=True, size=11, color=GREEN_HDR)

    hdr_row = s1_start + 1
    headers = ["Year", "Gas (GJ/PJ)", "Gas total (PJ)", "Diesel (GJ/PJ)",
               "Diesel total (PJ)", "Electricity (GJ/PJ)", "Electricity total (PJ)", "SEI (GJ/PJ)", "Total energy (PJ)"]
    for i, h in enumerate(headers, 1):
        c = ws1.cell(row=hdr_row, column=i, value=h)
        c.fill = hdr_fill(DARK_BLUE)
        c.font = cell_font(bold=True, color="FFFFFF", size=10)
        c.border = thin_border()
        c.alignment = Alignment(horizontal="center", wrap_text=True)
    ws1.row_dimensions[hdr_row].height = 30

    for i, yr in enumerate(YEARS):
        row = hdr_row + 1 + i
        demand = TOTAL_DEMAND[yr]
        g_coef = CONV_GAS[yr]
        r_coef = CONV_RFUEL[yr]
        e_coef = CONV_ELEC[yr]
        g_pj   = round(g_coef * demand / 1e6, 1)
        r_pj   = round(r_coef * demand / 1e6, 1)
        e_pj   = round(e_coef * demand / 1e6, 1)
        sei    = g_coef + r_coef + e_coef
        tot_pj = round(sei * demand / 1e6, 1)
        for j, v in enumerate([yr, g_coef, g_pj, r_coef, r_pj, e_coef, e_pj, sei, tot_pj], 1):
            c = ws1.cell(row=row, column=j, value=v)
            c.border = thin_border()
            c.alignment = Alignment(horizontal="right")
            c.font = cell_font(size=10)

    ck_row = hdr_row + 1 + len(YEARS)
    ws1.cell(row=ck_row, column=1, value=f"AES 2023-24 reference: {AES_FINAL_ENERGY:.0f} PJ").font = Font(bold=True, color="FFFFFF", size=10)
    ws1.cell(row=ck_row, column=8, value=f"Coverage: {pct(MODELLED_PJ, AES_FINAL_ENERGY)}").font = Font(bold=True, color="FFFFFF", size=10)
    ws1.cell(row=ck_row, column=9, value=f"~{AES_FINAL_ENERGY:.0f} PJ").font = Font(bold=True, color="FFFFFF", size=10)
    for col in range(1, 10):
        ws1.cell(row=ck_row, column=col).fill = hdr_fill(GREEN_HDR)
        ws1.cell(row=ck_row, column=col).border = thin_border()

    # Energy chart
    chart1 = BarChart()
    chart1.type = "col"
    chart1.grouping = "stacked"
    chart1.title = "Oil & Gas Extraction — Fuel by Type (PJ, Conventional State)"
    chart1.y_axis.title = "Energy (PJ)"
    chart1.x_axis.title = "Year"
    chart1.width = 16
    chart1.height = 10

    cats = Reference(ws1, min_col=1, min_row=hdr_row + 1, max_row=hdr_row + len(YEARS))
    chart1.add_data(Reference(ws1, min_col=3, min_row=hdr_row, max_row=hdr_row + len(YEARS)), titles_from_data=True)
    chart1.add_data(Reference(ws1, min_col=5, min_row=hdr_row, max_row=hdr_row + len(YEARS)), titles_from_data=True)
    chart1.add_data(Reference(ws1, min_col=7, min_row=hdr_row, max_row=hdr_row + len(YEARS)), titles_from_data=True)
    chart1.set_categories(cats)
    chart1.series[0].graphicalProperties.solidFill = MID_BLUE
    chart1.series[1].graphicalProperties.solidFill = ORANGE
    chart1.series[2].graphicalProperties.solidFill = "70AD47"
    ws1.add_chart(chart1, "A" + str(ck_row + 2))

    # ── Section 2: Emissions ──────────────────────────────────────────────────
    em_start = ck_row + 2 + 17
    ws1.cell(row=em_start, column=1, value="SECTION 2 — EMISSIONS (conventional state, tCO2e/PJ_gas × total production PJ)").font = Font(bold=True, size=11, color=GREEN_HDR)

    em_hdr = em_start + 1
    em_headers = ["Year", "Energy CO2e (tCO2e/PJ)", "Energy total (MtCO2e)",
                  "Fugitive CO2e (tCO2e/PJ)", "Fugitive total (MtCO2e)", "Total (MtCO2e)"]
    for i, h in enumerate(em_headers, 1):
        c = ws1.cell(row=em_hdr, column=i, value=h)
        c.fill = hdr_fill(DARK_BLUE)
        c.font = cell_font(bold=True, color="FFFFFF", size=10)
        c.border = thin_border()
        c.alignment = Alignment(horizontal="center", wrap_text=True)
    ws1.row_dimensions[em_hdr].height = 30

    for i, yr in enumerate(YEARS):
        row = em_hdr + 1 + i
        demand = TOTAL_DEMAND[yr]
        eco2e = CONV_E_CO2E[yr]
        pco2e = CONV_P_CO2E[yr]
        e_mt  = round(eco2e * demand / 1e6, 2)
        p_mt  = round(pco2e * demand / 1e6, 2)
        tot   = round(e_mt + p_mt, 2)
        for j, v in enumerate([yr, eco2e, e_mt, pco2e, p_mt, tot], 1):
            c = ws1.cell(row=row, column=j, value=v)
            c.border = thin_border()
            c.alignment = Alignment(horizontal="right")
            c.font = cell_font(size=10)

    em_ck_row = em_hdr + 1 + len(YEARS)
    ws1.cell(row=em_ck_row, column=1, value="NGGI 2023-24 reference").font = Font(bold=True, color="FFFFFF", size=10)
    ws1.cell(row=em_ck_row, column=3, value=f"~{NGGI_ENERGY_MT:.0f} MtCO2e (Cat 1A2)").font = Font(bold=True, color="FFFFFF", size=10)
    ws1.cell(row=em_ck_row, column=5, value=f"~{NGGI_FUGITIVE_MT:.0f} MtCO2e (Cat 1B2b)").font = Font(bold=True, color="FFFFFF", size=10)
    ws1.cell(row=em_ck_row, column=6, value=f"~{NGGI_ENERGY_MT + NGGI_FUGITIVE_MT:.0f} MtCO2e").font = Font(bold=True, color="FFFFFF", size=10)
    for col in range(1, 7):
        ws1.cell(row=em_ck_row, column=col).fill = hdr_fill(GREEN_HDR)
        ws1.cell(row=em_ck_row, column=col).border = thin_border()

    # AES-derived combustion note
    note_row = em_ck_row + 1
    ws1.cell(row=note_row, column=1, value=(
        f"Note: AES-derived combustion CO2e = {AES_DERIVED_CO2E_MT:.1f} MtCO2e "
        f"({GAS_2025:,}×51.4 + {RFUEL_2025:,}×69.9)/1,000 × {ANCHOR_PJ:,}. "
        f"Excess {AES_DERIVED_CO2E_MT - NGGI_ENERGY_MT:.1f} MtCO2e vs NGGI 1A2 = routine flaring/venting classified as fugitive (Cat 1B2b)."
    ))
    ws1.cell(row=note_row, column=1).font = Font(italic=True, size=9, color="595959")
    ws1.merge_cells(f"A{note_row}:I{note_row}")

    # Emissions chart
    chart2 = BarChart()
    chart2.type = "col"
    chart2.grouping = "stacked"
    chart2.title = "Oil & Gas Extraction — Emissions by Type (MtCO2e, Conventional State)"
    chart2.y_axis.title = "Emissions (MtCO2e)"
    chart2.x_axis.title = "Year"
    chart2.width = 16
    chart2.height = 10

    em_cats = Reference(ws1, min_col=1, min_row=em_hdr + 1, max_row=em_hdr + len(YEARS))
    chart2.add_data(Reference(ws1, min_col=3, min_row=em_hdr, max_row=em_hdr + len(YEARS)), titles_from_data=True)
    chart2.add_data(Reference(ws1, min_col=5, min_row=em_hdr, max_row=em_hdr + len(YEARS)), titles_from_data=True)
    chart2.set_categories(em_cats)
    chart2.series[0].graphicalProperties.solidFill = ORANGE
    chart2.series[1].graphicalProperties.solidFill = RED
    ws1.add_chart(chart2, "A" + str(em_ck_row + 3))

    ws1.freeze_panes = "B6"

    # ── Sheet 2: Calibration_Basis ────────────────────────────────────────────
    ws2 = wb.create_sheet("Calibration_Basis")
    ws2.sheet_properties.tabColor = MID_BLUE

    ws2["A1"] = "Oil & Gas Extraction — Calibration Basis and Parameters"
    ws2["A1"].font = Font(bold=True, size=13, color=DARK_BLUE)
    ws2.merge_cells("A1:E1")

    rows_cb = [
        ["Parameter", "Value", "Unit", "Source", "Notes"],
        ["Total gas production (anchor)", ANCHOR_PJ, "PJ_gas",
         "Geoscience Australia AECR 2025",
         "2023-24 total: LNG exports ~4,509 PJ (74%) + domestic ~1,591 PJ (26%)"],
        ["AES final energy (oil & gas extraction)", AES_FINAL_ENERGY, "PJ",
         "AES 2023-24 Table F — DCCEEW",
         "ANZSIC 0700 final energy. Excludes LNG liquefaction own-use (classified as energy transformation)."],
        ["Previous model energy (gas_mining)", PREV_ENERGY_PJ, "PJ",
         "Old gas_mining generator (AES estimate)",
         f"Overcount: included LNG liquefaction own-use ~{PREV_ENERGY_PJ - AES_FINAL_ENERGY:.0f} PJ."],
        ["Gas own-use coefficient (revised)", GAS_2025, "GJ/PJ_gas",
         "AES 2023-24 Table F / Geoscience Australia",
         "380 PJ / 6,100 PJ. Upstream only (well compression + gathering + processing)."],
        ["Gas own-use coefficient (previous)", PREV_GAS_COEF, "GJ/PJ_gas",
         "Old estimate",
         "Incorrectly included LNG liquefaction own-use (~316 PJ from Geoscience Australia 7% LNG factor)."],
        ["Diesel coefficient", RFUEL_2025, "GJ/PJ_gas",
         "AES 2023-24 Table F",
         "60 PJ / 6,100 PJ. Drilling rigs, earthworks, auxiliary plant."],
        ["Electricity coefficient", ELEC_2025, "GJ/PJ_gas",
         "AES 2023-24 Table F",
         "31 PJ / 6,100 PJ. Fixed facility loads."],
        ["Total SEI (revised)", SEI_2025, "GJ/PJ_gas",
         "Derived",
         f"Sense-check: {SEI_2025:,} × {ANCHOR_PJ:,} / 1e6 = {MODELLED_PJ:.1f} PJ ≈ AES {AES_FINAL_ENERGY:.0f} PJ ✓"],
        ["Energy CO2e (revised)", ENERGY_CO2E_2025, "tCO2e/PJ_gas",
         "NGGI 2023-24 Cat 1A2 (calibrated)",
         f"22 MtCO2e / {ANCHOR_PJ:,} PJ = 3,607 → 3,610 tCO2e/PJ. NGGI-direct calibration."],
        ["Energy CO2e (AES-derived, pre-NGGI)", round(AES_DERIVED_CO2E_MT * 1e6 / ANCHOR_PJ), "tCO2e/PJ_gas",
         "NGA 2024 EFs: gas 51.4, diesel 69.9 kgCO2e/GJ",
         f"({GAS_2025:,}×51.4 + {RFUEL_2025:,}×69.9)/1,000 = 3,887 → {AES_DERIVED_CO2E_MT:.1f} MtCO2e. Gap {AES_DERIVED_CO2E_MT - NGGI_ENERGY_MT:.1f} MtCO2e = flaring/venting in fugitive."],
        ["Energy CO2e (previous model)", PREV_ENERGY_CO2E, "tCO2e/PJ_gas",
         "Old generator (NGA 2024 EFs × wrong gas coeff)",
         f"Overcount: {PREV_ENERGY_CO2E * ANCHOR_PJ / 1e6:.1f} MtCO2e vs NGGI 22.0 MtCO2e."],
        ["Fugitive CO2e (revised)", PROCESS_CO2E_2025, "tCO2e/PJ_gas",
         "NGGI 2023-24 Cat 1B2b",
         f"22 MtCO2e / {ANCHOR_PJ:,} PJ = 3,607 → 3,610. Includes flaring/venting gas reclassified from energy."],
        ["Fugitive CO2e (previous model)", PREV_PROCESS_CO2E, "tCO2e/PJ_gas",
         "NGGI 2022-23 Cat 1B2b",
         f"{PREV_PROCESS_CO2E * ANCHOR_PJ / 1e6:.1f} MtCO2e. Used older inventory year; undercount by 2 MtCO2e."],
        ["Demand trajectory", "−1.0%/yr", "compound",
         "AEMO ISP 2024 Step Change",
         "Global LNG demand peaking post-2030; domestic gas declining as electrification proceeds."],
    ]

    for i, row_data in enumerate(rows_cb):
        for j, val in enumerate(row_data, 1):
            c = ws2.cell(row=3 + i, column=j, value=val)
            c.border = thin_border()
            if i == 0:
                c.fill = hdr_fill(DARK_BLUE)
                c.font = cell_font(bold=True, color="FFFFFF", size=10)
                c.alignment = Alignment(horizontal="center", wrap_text=True)
            else:
                c.font = cell_font(size=9)
                c.alignment = Alignment(wrap_text=True, vertical="top")
        if i > 0:
            ws2.row_dimensions[3 + i].height = 36

    set_col_widths(ws2, {"A": 36, "B": 14, "C": 16, "D": 34, "E": 60})

    # ── Sheet 3: State_Assumptions ────────────────────────────────────────────
    ws3 = wb.create_sheet("State_Assumptions")
    ws3.sheet_properties.tabColor = AMBER

    ws3["A1"] = "Oil & Gas Extraction — Technology State Assumptions (2025, Revised)"
    ws3["A1"].font = Font(bold=True, size=13, color=DARK_BLUE)
    ws3.merge_cells("A1:K1")

    state_hdr = ["State ID", "Label", "Gas (GJ/PJ)", "Diesel (GJ/PJ)", "Elec (GJ/PJ)",
                 "SEI (GJ/PJ)", "EEI", "Energy CO2e (tCO2e/PJ)", "Fugitive CO2e (tCO2e/PJ)",
                 "Max share 2025", "Confidence"]
    state_data_rows = [
        ["oil_and_gas_extraction__conventional",           "Conventional (gas turbine + diesel + grid)", 62_300,  9_800,  5_100, 77_200, 1.000, 3_610, 3_610, "1.00", "Medium"],
        ["oil_and_gas_extraction__electric_compression",   "Electric motor drive compression",           23_000,  9_800, 17_500, 50_300, 0.651, 1_730, 3_610, "0.01", "Low-Medium"],
        ["oil_and_gas_extraction__ldar_no_flare",          "LDAR and zero routine flaring",              62_300,  9_800,  5_200, 77_300, 1.001, 3_610, 1_805, "0.02", "Medium"],
        ["oil_and_gas_extraction__renewable_diesel",       "Remote wellsite solar-diesel hybrid",        62_300,  4_900,  6_600, 73_800, 0.956, 3_290, 3_610, "0.01", "Medium"],
        ["oil_and_gas_extraction__integrated_low_emission","Integrated low-emission",                    19_000,  4_900, 18_500, 42_400, 0.549, 1_230, 1_805, "0.00", "Exploratory"],
    ]

    for j, h in enumerate(state_hdr, 1):
        c = ws3.cell(row=3, column=j, value=h)
        c.fill = hdr_fill(DARK_BLUE)
        c.font = cell_font(bold=True, color="FFFFFF", size=10)
        c.border = thin_border()
        c.alignment = Alignment(horizontal="center", wrap_text=True)
    ws3.row_dimensions[3].height = 36

    for i, row_data in enumerate(state_data_rows):
        for j, val in enumerate(row_data, 1):
            c = ws3.cell(row=4 + i, column=j, value=val)
            c.border = thin_border()
            c.font = cell_font(size=9)
            c.alignment = Alignment(horizontal="right" if j > 2 else "left", wrap_text=j <= 2)
        ws3.row_dimensions[4 + i].height = 24

    set_col_widths(ws3, {"A": 42, "B": 36, "C": 13, "D": 13, "E": 13,
                         "F": 13, "G": 8, "H": 20, "I": 22, "J": 12, "K": 14})

    # ── Sheet 4: Total_Demand ─────────────────────────────────────────────────
    ws4 = wb.create_sheet("Total_Demand")
    ws4.sheet_properties.tabColor = MID_BLUE

    ws4["A1"] = "Oil & Gas Extraction — Demand Trajectory (declining__oil_and_gas_extraction_total, −1.0%/yr)"
    ws4["A1"].font = Font(bold=True, size=13, color=DARK_BLUE)
    ws4.merge_cells("A1:D1")

    dem_hdr = ["Year", "Total demand (PJ_gas)", "Index (2025=1.000)", "Notes"]
    for j, h in enumerate(dem_hdr, 1):
        c = ws4.cell(row=3, column=j, value=h)
        c.fill = hdr_fill(DARK_BLUE)
        c.font = cell_font(bold=True, color="FFFFFF", size=10)
        c.border = thin_border()
        c.alignment = Alignment(horizontal="center", wrap_text=True)

    for i, yr in enumerate(YEARS):
        row = 4 + i
        demand = TOTAL_DEMAND[yr]
        idx    = round(demand / ANCHOR_PJ, 4)
        note   = "Anchor (Geoscience Australia AECR 2025)" if yr == 2025 else ""
        for j, val in enumerate([yr, demand, idx, note], 1):
            c = ws4.cell(row=row, column=j, value=val)
            c.border = thin_border()
            c.font = cell_font(size=10)
            c.alignment = Alignment(horizontal="right" if j < 4 else "left")

    set_col_widths(ws4, {"A": 10, "B": 22, "C": 18, "D": 46})

    # ── Sheet 5: Sources ──────────────────────────────────────────────────────
    ws5 = wb.create_sheet("Sources")
    ws5.sheet_properties.tabColor = "595959"

    ws5["A1"] = "Oil & Gas Extraction — Sources and Assumptions Register"
    ws5["A1"].font = Font(bold=True, size=13, color=DARK_BLUE)
    ws5.merge_cells("A1:D1")

    src_hdr = ["ID", "Type", "Reference", "Used for"]
    sources = [
        ["S001", "Source",
         "Australian Energy Statistics (AES) 2023-24, Table F — DCCEEW",
         "Final energy 471 PJ for oil and gas extraction (ANZSIC 0700). Excludes LNG liquefaction own-use."],
        ["S002", "Source",
         "National Greenhouse Gas Inventory (NGGI) 2023-24 — DCCEEW",
         "Cat 1A2 energy emissions 22 MtCO2e; Cat 1B2b fugitive 22 MtCO2e. Both calibration targets."],
        ["S022", "Source",
         "Geoscience Australia — Australian Energy Commodity Report (AECR) 2025",
         "Total gas production anchor 6,100 PJ; LNG exports 4,509 PJ; domestic supply 1,591 PJ."],
        ["S023", "Source",
         "National Greenhouse and Energy Reporting (NGER) data 2022-23 — CER",
         "Facility-level energy and emissions for upstream operators (Santos, Woodside, Beach, Origin)."],
        ["S024", "Source",
         "NGA 2024 — National Greenhouse Accounts Emission Factors — DCCEEW",
         "EF natural gas combustion 51.4 kgCO2e/GJ; diesel 69.9 kgCO2e/GJ (AR5 GWP100)."],
        ["S025", "Source",
         "IEA Natural Gas Decarbonisation 2022; IEA Global Methane Tracker 2024",
         "LDAR methane reduction potential; electric compression technology costs; pathway benchmarks."],
        ["S026", "Source",
         "AEMO Integrated System Plan (ISP) 2024 — Step Change scenario",
         "Gas demand trajectory; domestic gas demand decline rate as electrification proceeds."],
        ["A002", "Assumption",
         "AES final energy (471 PJ) covers only upstream oil and gas extraction (ANZSIC 0700) final energy",
         "LNG liquefaction own-use (~316 PJ) is classified in AES as 'energy transformation', not 'final energy consumption'. This is the main source of the 595→471 PJ correction."],
        ["A003", "Assumption",
         "Scope 1 boundary: direct combustion (Cat 1A2) and fugitive (Cat 1B2b) only",
         "Routine flaring/venting gas combustion is classified as fugitive (1B2b) not energy (1A2) in NGGI. The 1.7 MtCO2e gap between AES-derived combustion and NGGI 1A2 is explained by this reclassification."],
        ["A010", "Assumption",
         "Per-PJ_gas coefficients are national averages combining onshore CSG, conventional, and LNG",
         "Phase 2: disaggregate into QLD CSG, WA/NT onshore conventional, and offshore LNG sub-families. National average used in Phase 1 only."],
    ]

    for j, h in enumerate(src_hdr, 1):
        c = ws5.cell(row=3, column=j, value=h)
        c.fill = hdr_fill(DARK_BLUE)
        c.font = cell_font(bold=True, color="FFFFFF", size=10)
        c.border = thin_border()
        c.alignment = Alignment(horizontal="center")

    for i, row_data in enumerate(sources):
        for j, val in enumerate(row_data, 1):
            c = ws5.cell(row=4 + i, column=j, value=val)
            c.border = thin_border()
            c.font = cell_font(size=9)
            c.alignment = Alignment(wrap_text=True, vertical="top")
        ws5.row_dimensions[4 + i].height = 52

    set_col_widths(ws5, {"A": 8, "B": 12, "C": 64, "D": 64})

    path = OUT_DIR / "oil_and_gas_extraction_calibration.xlsx"
    wb.save(path)
    print(f"  Written: {path}")


# ═══════════════════════════════════════════════════════════════════════════════
#  WORD DOCUMENT
# ═══════════════════════════════════════════════════════════════════════════════
def set_cell_bg(cell, hex_color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shd)

def table_hdr_cell(cell, text, bg_hex=DARK_BLUE):
    cell.text = text
    run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(text)
    if not cell.paragraphs[0].runs:
        cell.text = ""
        run = cell.paragraphs[0].add_run(text)
    else:
        run = cell.paragraphs[0].runs[0]
        run.text = text
    run.font.size = Pt(8)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    set_cell_bg(cell, bg_hex)

def table_data_cell(cell, text, bold=False, bg_hex=None):
    cell.text = str(text)
    if cell.paragraphs[0].runs:
        run = cell.paragraphs[0].runs[0]
        run.font.size = Pt(8)
        run.font.bold = bold
        if bg_hex:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    if bg_hex:
        set_cell_bg(cell, bg_hex)


def build_word():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    title = doc.add_heading("Oil and Gas Extraction — Calibration Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        f"Phase 1 | Total Australian oil and gas production | Generated: {datetime.date.today().isoformat()}"
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    # ── Section 1: Scope ─────────────────────────────────────────────────────
    doc.add_heading("1. Scope and Demand Anchor", level=1)
    doc.add_paragraph(
        f"This family models total Australian oil and gas extraction (gas-production-weighted) "
        f"as a single demand unit. Demand anchor: {ANCHOR_PJ:,} PJ_gas (2025; "
        f"Geoscience Australia AECR 2025), covering LNG exports (~4,509 PJ, 74%) and "
        f"domestic/pipeline supply (~1,591 PJ, 26%). Service output unit: PJ of gas produced "
        f"(PJ_gas). Demand trajectory: declining__oil_and_gas_extraction_total (−1.0%/yr compound, "
        f"AEMO ISP 2024 Step Change scenario, reflecting global LNG demand peaking post-2030)."
    )
    doc.add_paragraph(
        f"Per-PJ_gas energy coefficients are national averages calibrated to AES 2023-24 "
        f"Table F (oil and gas extraction, ANZSIC 0700, {AES_FINAL_ENERGY:.0f} PJ final energy). "
        f"The sector mixes onshore CSG (QLD), conventional onshore (NT, WA), and offshore "
        f"LNG facilities (NW Shelf, Browse, Ichthys, Prelude, Darwin LNG). "
        f"Phase 2 disaggregation by sub-sector is recommended."
    )

    # ── Section 2: Calibration Discrepancies ─────────────────────────────────
    doc.add_heading("2. Calibration Revision — Discrepancies vs Previous Model", level=1)
    doc.add_paragraph(
        "This family replaces the previous gas_mining family. Three material discrepancies "
        "between the previous model and official AES/NGGI data were identified and corrected:"
    )

    # Discrepancy 1 — Energy
    doc.add_heading("2.1 Final Energy: Previous 595 PJ → Revised 471 PJ", level=2)
    doc.add_paragraph(
        f"Source of truth: AES 2023-24 Table F (DCCEEW). "
        f"The previous model estimated gas own-use = 82,000 GJ/PJ based on Geoscience Australia's "
        f"estimate that LNG facilities consume ~7% of throughput as fuel (~316 PJ from 4,509 PJ "
        f"LNG exports) plus upstream compression (~184 PJ), totalling ~500 PJ and an SEI of "
        f"97,600 GJ/PJ → 595 PJ."
    )
    doc.add_paragraph(
        f"However, AES 2023-24 'final energy consumption' for oil and gas extraction (ANZSIC 0700) "
        f"= {AES_FINAL_ENERGY:.0f} PJ. This is a lower number because AES classifies LNG "
        f"liquefaction own-use under the 'energy transformation' sector, not 'final energy "
        f"consumption'. The correct AES-compatible gas coefficient is 62,300 GJ/PJ "
        f"(380 PJ upstream / 6,100 PJ), which covers only wellhead compression, gathering, "
        f"and processing — not LNG liquefaction. "
        f"Revised SEI: 77,200 GJ/PJ × 6,100 PJ / 1e6 = {MODELLED_PJ:.1f} PJ ({pct(MODELLED_PJ, AES_FINAL_ENERGY)} coverage) ✓"
    )

    # Discrepancy 2 — Energy CO2e
    doc.add_heading("2.2 Energy CO2e: Previous 29.5 MtCO2e → Revised 22.0 MtCO2e", level=2)
    doc.add_paragraph(
        "Source of truth: NGGI 2023-24 Category 1A2 (stationary energy — mining) = 22 MtCO2e. "
        "Two causes:"
    )
    p = doc.add_paragraph(style="List Bullet")
    p.add_run("(a) Wrong gas coefficient: ").bold = True
    p.add_run(
        f"The previous 82,000 GJ/PJ coefficient was the dominant driver. "
        f"With the corrected coefficient (62,300 GJ/PJ), "
        f"the NGA 2024 emission factor calculation gives: "
        f"(62,300 × 51.4 + 9,800 × 69.9) / 1,000 = 3,887 tCO2e/PJ "
        f"→ {AES_DERIVED_CO2E_MT:.1f} MtCO2e."
    )
    p = doc.add_paragraph(style="List Bullet")
    p.add_run("(b) Flaring/venting reclassification: ").bold = True
    p.add_run(
        f"The remaining {AES_DERIVED_CO2E_MT - NGGI_ENERGY_MT:.1f} MtCO2e gap "
        f"({AES_DERIVED_CO2E_MT:.1f} MtCO2e AES-derived vs NGGI 1A2 22.0 MtCO2e) arises because "
        f"gas combusted through routine flaring and venting at oil and gas facilities is classified "
        f"under NGGI Category 1B2b (Fugitive: oil and gas), not Category 1A2 (Energy: stationary "
        f"combustion). This {AES_DERIVED_CO2E_MT - NGGI_ENERGY_MT:.1f} MtCO2e is correctly captured "
        f"in the fugitive (process) CO2e coefficient. The energy CO2e is therefore calibrated "
        f"directly to NGGI 1A2 = {ENERGY_CO2E_2025:,} tCO2e/PJ ({ENERGY_CO2E_2025 * ANCHOR_PJ / 1e6:.1f} MtCO2e)."
    )

    # Discrepancy 3 — Fugitive
    doc.add_heading("2.3 Fugitive CO2e: Previous 20.0 MtCO2e → Revised 22.0 MtCO2e", level=2)
    doc.add_paragraph(
        "Source of truth: NGGI 2023-24 Category 1B2b (oil and gas fugitive) = 22 MtCO2e. "
        "The previous model used NGGI 2022-23 data (~20 MtCO2e). The 2 MtCO2e increase in "
        "NGGI 2023-24 reflects two factors: (1) improved measurement methodology — mandatory "
        "direct-measurement reporting under NGER for large facilities captures previously "
        "underestimated wellsite methane; (2) real growth in LNG sector fugitive emissions "
        "as production ramped up. The revised coefficient is 3,610 tCO2e/PJ "
        f"(22,000,000 / 6,100 = 3,607, rounded to 3,610). "
        f"Note: this includes the {AES_DERIVED_CO2E_MT - NGGI_ENERGY_MT:.1f} MtCO2e of flaring/venting "
        "gas reclassified from energy (1A2) to fugitive (1B2b)."
    )

    doc.add_paragraph("")

    # ── Section 3: Technology States ─────────────────────────────────────────
    doc.add_heading("3. Technology States (2025 Summary)", level=1)
    doc.add_paragraph(
        "Five technology states are defined covering the full range from conventional "
        "gas turbine operations to fully integrated low-emission production."
    )

    tbl = doc.add_table(rows=1, cols=7)
    tbl.style = "Table Grid"
    hdr_vals = ["State", "Gas (GJ/PJ)", "Diesel (GJ/PJ)", "Elec (GJ/PJ)",
                "SEI (GJ/PJ)", "Energy CO2e (tCO2e/PJ)", "Fugitive CO2e (tCO2e/PJ)"]
    for i, h in enumerate(hdr_vals):
        table_hdr_cell(tbl.rows[0].cells[i], h)

    state_rows_doc = [
        ("Conventional",           62_300,  9_800,  5_100, 77_200, 3_610, 3_610),
        ("Electric compression",   23_000,  9_800, 17_500, 50_300, 1_730, 3_610),
        ("LDAR / no-flare",        62_300,  9_800,  5_200, 77_300, 3_610, 1_805),
        ("Renewable diesel",       62_300,  4_900,  6_600, 73_800, 3_290, 3_610),
        ("Integrated low-emission", 19_000,  4_900, 18_500, 42_400, 1_230, 1_805),
    ]
    for sr in state_rows_doc:
        row_cells = tbl.add_row().cells
        for i, v in enumerate(sr):
            row_cells[i].text = f"{v:,}" if isinstance(v, int) else str(v)
            if row_cells[i].paragraphs[0].runs:
                row_cells[i].paragraphs[0].runs[0].font.size = Pt(8)

    doc.add_paragraph("")

    # ── Section 4: Energy Calibration ────────────────────────────────────────
    doc.add_heading("4. Energy Calibration (Conventional State)", level=1)
    doc.add_paragraph(
        f"The incumbent conventional state is calibrated to AES 2023-24 Table F. "
        f"Gas own-use ({GAS_2025:,} GJ/PJ) covers upstream wellhead compression, "
        f"gathering pipelines, and gas processing (380 PJ / 6,100 PJ). "
        f"Diesel ({RFUEL_2025:,} GJ/PJ = {RFUEL_2025 * ANCHOR_PJ / 1e6:.0f} PJ) covers "
        f"remote drilling rigs and auxiliary plant. "
        f"Electricity ({ELEC_2025:,} GJ/PJ = {ELEC_2025 * ANCHOR_PJ / 1e6:.0f} PJ) covers "
        f"fixed facility loads. LNG liquefaction own-use is excluded from AES final energy "
        f"(classified as 'energy transformation') and therefore excluded from per-PJ coefficients."
    )

    doc.add_heading("4.1 Fuel Consumption Sense-Check", level=2)
    tbl2 = doc.add_table(rows=1, cols=5)
    tbl2.style = "Table Grid"
    for i, h in enumerate(["Fuel", "Coeff. (GJ/PJ)", "Total (PJ)", "AES reference (PJ)", "Coverage"]):
        table_hdr_cell(tbl2.rows[0].cells[i], h)

    fuel_rows = [
        ("Gas own-use",  GAS_2025,  round(GAS_2025  * ANCHOR_PJ / 1e6, 0), "~380", pct(GAS_2025 * ANCHOR_PJ / 1e6, 380)),
        ("Diesel",       RFUEL_2025, round(RFUEL_2025 * ANCHOR_PJ / 1e6, 0), "~60",  pct(RFUEL_2025 * ANCHOR_PJ / 1e6, 60)),
        ("Electricity",  ELEC_2025,  round(ELEC_2025  * ANCHOR_PJ / 1e6, 0), "~31",  pct(ELEC_2025 * ANCHOR_PJ / 1e6, 31)),
        ("Total",        SEI_2025,   round(SEI_2025   * ANCHOR_PJ / 1e6, 0), f"~{AES_FINAL_ENERGY:.0f}", pct(MODELLED_PJ, AES_FINAL_ENERGY)),
    ]
    for fr in fuel_rows:
        row_cells = tbl2.add_row().cells
        for i, v in enumerate(fr):
            row_cells[i].text = f"{v:,.0f}" if isinstance(v, float) else (f"{v:,}" if isinstance(v, int) else str(v))
            if row_cells[i].paragraphs[0].runs:
                row_cells[i].paragraphs[0].runs[0].font.size = Pt(8)

    # ── Section 5: Emissions Calibration ─────────────────────────────────────
    doc.add_heading("5. Emissions Calibration", level=1)

    doc.add_heading("5.1 Energy Emissions (Combustion, NGGI Cat 1A2)", level=2)
    doc.add_paragraph(
        f"NGA 2024 emission factors (DCCEEW, AR5 GWP100): "
        f"natural gas = 51.4 kgCO2e/GJ; diesel = 69.9 kgCO2e/GJ. "
        f"AES-derived combustion CO2e: "
        f"({GAS_2025:,} × 51.4 + {RFUEL_2025:,} × 69.9) / 1,000 = 3,887 tCO2e/PJ "
        f"→ {AES_DERIVED_CO2E_MT:.1f} MtCO2e. "
        f"NGGI 2023-24 Cat 1A2 = 22.0 MtCO2e. "
        f"The {AES_DERIVED_CO2E_MT - NGGI_ENERGY_MT:.1f} MtCO2e gap is because routine "
        f"flaring and venting gas combustion is classified as fugitive (Cat 1B2b) in NGGI, "
        f"not as energy (Cat 1A2). Model energy CO2e is calibrated directly to NGGI 1A2: "
        f"{ENERGY_CO2E_2025:,} tCO2e/PJ × {ANCHOR_PJ:,} PJ = {ENERGY_CO2E_2025 * ANCHOR_PJ / 1e6:.1f} MtCO2e ✓"
    )

    doc.add_heading("5.2 Fugitive Emissions (Process, NGGI Cat 1B2b)", level=2)
    doc.add_paragraph(
        f"NGGI 2023-24 Category 1B2b (oil and gas fugitive): 22.0 MtCO2e. "
        f"Per PJ_gas: 22,000,000 / {ANCHOR_PJ:,} = 3,607 → modelled as {PROCESS_CO2E_2025:,} tCO2e/PJ_gas. "
        f"Sense-check: {PROCESS_CO2E_2025:,} × {ANCHOR_PJ:,} / 1e6 = {PROCESS_CO2E_2025 * ANCHOR_PJ / 1e6:.1f} MtCO2e "
        f"≈ NGGI {NGGI_FUGITIVE_MT:.0f} MtCO2e ✓. "
        f"Fugitive sources include: wellhead venting, compressor seal gas leaks, "
        f"pipeline fugitives, LNG processing purge streams, and routine flaring/venting "
        f"gas (reclassified from Cat 1A2 by NGGI methodology). "
        f"Previous model used NGGI 2022-23 ({PREV_PROCESS_CO2E:,} tCO2e/PJ = "
        f"{PREV_PROCESS_CO2E * ANCHOR_PJ / 1e6:.1f} MtCO2e), understating by 2 MtCO2e."
    )

    doc.add_heading("5.3 Emissions Sense-Check Table", level=2)
    tbl3 = doc.add_table(rows=1, cols=5)
    tbl3.style = "Table Grid"
    for i, h in enumerate(["Year", "Energy CO2e (tCO2e/PJ)", "Energy total (MtCO2e)",
                             "Fugitive (tCO2e/PJ)", "Fugitive total (MtCO2e)"]):
        table_hdr_cell(tbl3.rows[0].cells[i], h)

    for yr in YEARS:
        demand = TOTAL_DEMAND[yr]
        eco2e  = CONV_E_CO2E[yr]
        pco2e  = CONV_P_CO2E[yr]
        row_cells = tbl3.add_row().cells
        for i, v in enumerate([yr,
                                f"{eco2e:,}",
                                f"{eco2e * demand / 1e6:.2f}",
                                f"{pco2e:,}",
                                f"{pco2e * demand / 1e6:.2f}"]):
            row_cells[i].text = str(v)
            if row_cells[i].paragraphs[0].runs:
                row_cells[i].paragraphs[0].runs[0].font.size = Pt(8)

    ck_cells = tbl3.add_row().cells
    for i, v in enumerate(["NGGI 2023-24", "—",
                             f"{NGGI_ENERGY_MT:.0f} MtCO2e (Cat 1A2) ✓",
                             "—",
                             f"{NGGI_FUGITIVE_MT:.0f} MtCO2e (Cat 1B2b) ✓"]):
        table_data_cell(ck_cells[i], v, bold=True, bg_hex=GREEN_HDR)

    doc.add_paragraph("")

    # ── Section 6: Demand Trajectory ─────────────────────────────────────────
    doc.add_heading("6. Demand Trajectory (2025–2050)", level=1)
    doc.add_paragraph(
        "Total Australian gas production is modelled declining at −1.0%/yr compound "
        "(curve: declining__oil_and_gas_extraction_total), reflecting AEMO ISP 2024 "
        "Step Change scenario projections of global LNG demand peaking ~2030 and "
        "declining domestic gas demand as building and industrial electrification proceeds."
    )

    tbl4 = doc.add_table(rows=1, cols=3)
    tbl4.style = "Table Grid"
    for i, h in enumerate(["Year", "Total demand (PJ_gas)", "Index (2025=1.000)"]):
        table_hdr_cell(tbl4.rows[0].cells[i], h)

    for yr in YEARS:
        demand = TOTAL_DEMAND[yr]
        row_cells = tbl4.add_row().cells
        for i, v in enumerate([yr, f"{demand:,}", f"{demand/ANCHOR_PJ:.4f}"]):
            row_cells[i].text = str(v)
            if row_cells[i].paragraphs[0].runs:
                row_cells[i].paragraphs[0].runs[0].font.size = Pt(8)

    doc.add_paragraph("")

    # ── Section 7: Sources ────────────────────────────────────────────────────
    doc.add_heading("7. Sources and Assumptions", level=1)
    sources_doc = [
        ("S001", "AES 2023-24 Table F", "DCCEEW",
         f"Final energy {AES_FINAL_ENERGY:.0f} PJ for oil and gas extraction (ANZSIC 0700). Excludes LNG liquefaction own-use."),
        ("S002", "NGGI 2023-24 Cat 1A2 and 1B2b", "DCCEEW",
         f"Energy emissions {NGGI_ENERGY_MT:.0f} MtCO2e (Cat 1A2); fugitive {NGGI_FUGITIVE_MT:.0f} MtCO2e (Cat 1B2b). Both calibration targets."),
        ("S022", "AECR 2025", "Geoscience Australia",
         f"Total gas production {ANCHOR_PJ:,} PJ; LNG exports 4,509 PJ; domestic 1,591 PJ."),
        ("S023", "NGER data 2022-23", "Clean Energy Regulator",
         "Facility-level energy and emissions, upstream gas operators (Santos, Woodside, Beach, Origin)."),
        ("S024", "NGA 2024 Emission Factors", "DCCEEW",
         "Gas 51.4, diesel 69.9 kgCO2e/GJ (AR5 GWP100)."),
        ("S025", "IEA Natural Gas Decarbonisation 2022; Global Methane Tracker 2024", "IEA",
         "LDAR potential, electric compression costs, methane abatement benchmarks."),
        ("S026", "AEMO ISP 2024 Step Change", "AEMO",
         "Gas demand trajectory −1.0%/yr; domestic decline assumptions."),
        ("A002", "AES boundary assumption", "—",
         "AES final energy (471 PJ) covers upstream ANZSIC 0700 only. LNG liquefaction classified under 'energy transformation'. This boundary correction accounts for the 595→471 PJ change."),
        ("A003", "NGGI 1A2 / 1B2b boundary", "—",
         f"Flaring/venting gas combustion is 1B2b (fugitive), not 1A2 (energy). The {AES_DERIVED_CO2E_MT - NGGI_ENERGY_MT:.1f} MtCO2e gap between AES-derived combustion and NGGI 1A2 is explained by this."),
        ("A010", "National average coefficients", "—",
         "Phase 1: national average mixing onshore CSG, conventional, and offshore LNG. Phase 2: disaggregate sub-families."),
    ]

    tbl5 = doc.add_table(rows=1, cols=4)
    tbl5.style = "Table Grid"
    for i, h in enumerate(["ID", "Reference", "Publisher", "Used for"]):
        table_hdr_cell(tbl5.rows[0].cells[i], h)

    for s in sources_doc:
        row_cells = tbl5.add_row().cells
        for i, v in enumerate(s):
            row_cells[i].text = str(v)
            if row_cells[i].paragraphs[0].runs:
                row_cells[i].paragraphs[0].runs[0].font.size = Pt(8)

    path = OUT_DIR / "oil_and_gas_extraction_calibration_report.docx"
    doc.save(path)
    print(f"  Written: {path}")


if __name__ == "__main__":
    print(f"\nGenerating oil_and_gas_extraction calibration documents...\n")
    build_excel()
    build_word()
    print("\nDone.")
